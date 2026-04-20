# ==========================================
# 🚀 SAVIA FastAPI Backend
# ==========================================
"""
FastAPI backend for the SAVIA Next.js frontend.
Replaces Flask api_server.py with modern async endpoints.
"""
import math
import os
import jwt
import bcrypt
import logging
import auth
import pandas as pd
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any

from fpdf import FPDF
from fpdf.enums import XPos, YPos
from fastapi import FastAPI, Depends, HTTPException, Query, Header, status, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel

from db_engine import (
    init_db, get_db, read_sql,
    lire_equipements, ajouter_equipement, modifier_equipement, supprimer_equipement,
    lire_interventions, ajouter_intervention, update_intervention_statut, cloturer_intervention,
    lire_pieces, ajouter_piece, modifier_piece, supprimer_piece,
    lire_notifications_pieces, compter_notifications_non_lues, ajouter_notification_piece,
    marquer_notification_lue, marquer_notification_traitee, notifications_rupture_pour_piece,
    lire_contrats, ajouter_contrat, modifier_contrat, supprimer_contrat,
    lire_conformite, ajouter_conformite, supprimer_conformite,
    lire_planning, ajouter_planning, update_planning_statut, supprimer_planning,
    lire_techniciens, ajouter_technicien, update_technicien, supprimer_technicien,
    lire_base,
    lire_audit, log_audit,
    get_config, set_config,
    lire_demandes_intervention,
)

logger = logging.getLogger("savia-api")
logging.basicConfig(level=logging.INFO)

# ---- Config ----
JWT_SECRET = os.getenv("JWT_SECRET", "sic-terrain-secret-2026")
JWT_EXPIRY_HOURS = 72
security = HTTPBearer(auto_error=False)

# ---- App ----
app = FastAPI(
    title="SAVIA API",
    description="Backend API for SAVIA — Superviseur Intelligent Clinique",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup():
    init_db()
    auth.creer_admin_defaut()
    # Migration: colonnes fiche signée
    try:
        with get_db() as conn:
            conn.execute("ALTER TABLE interventions ADD COLUMN IF NOT EXISTS fiche_photo_nom TEXT DEFAULT ''")
            conn.execute("ALTER TABLE interventions ADD COLUMN IF NOT EXISTS fiche_photo_data BYTEA")
            conn.execute("ALTER TABLE interventions ADD COLUMN IF NOT EXISTS fiche_validation TEXT DEFAULT 'En attente'")
        logger.info("✅ Migration fiche_photo: colonnes OK")
    except Exception as e:
        logger.info(f"Migration fiche_photo (déjà faite ou erreur): {e}")
    _start_garantie_daemon()
    logger.info("✅ SAVIA FastAPI started — DB initialized")


# ==========================================
# HELPERS
# ==========================================


def check_garantie_expiry():
    """
    Vérifie les garanties équipements qui expirent dans les 30 prochains jours
    et envoie une notification Telegram pour chaque équipement concerné.
    """
    from datetime import date, timedelta
    try:
        df = lire_equipements()
        if df is None or df.empty:
            return []
        today = date.today()
        alert_limit = today + timedelta(days=30)
        alerts = []
        for _, row in df.iterrows():
            debut_str = str(row.get('garantie_debut', '') or '').strip()
            duree = int(row.get('garantie_duree', 0) or 0)
            if not debut_str or not duree:
                continue
            try:
                debut = date.fromisoformat(debut_str[:10])
                # Add years without dateutil
                try:
                    fin = debut.replace(year=debut.year + duree)
                except ValueError:  # Feb 29 edge case
                    fin = debut.replace(year=debut.year + duree, day=28)
                if today <= fin <= alert_limit:
                    alerts.append({
                        'nom': row.get('Nom') or row.get('nom', '?'),
                        'client': row.get('Client') or row.get('client', '?'),
                        'fin': fin.strftime('%d/%m/%Y'),
                        'jours': (fin - today).days,
                    })
            except Exception:
                continue
        if alerts:
            lines = '\n'.join(
                f"  • <b>{a['nom']}</b> ({a['client']}) — expire le {a['fin']} ({a['jours']}j)"
                for a in alerts
            )
            msg = (
                f"⚠️ <b>Garanties expirant bientôt</b>\n\n"
                f"{lines}\n\n"
                f"📅 Vérification SAVIA — {today.strftime('%d/%m/%Y')}"
            )
            _send_telegram(msg)
            logger.info(f"Garantie check: {len(alerts)} alerte(s) envoyée(s)")
        return alerts
    except Exception as e:
        logger.error(f"Garantie expiry check error: {e}")
        return []


def _start_garantie_daemon():
    """Lance un thread démon qui vérifie garanties + contrats toutes les 24h."""
    import threading, time
    def _run():
        time.sleep(30)
        while True:
            try:
                check_garantie_expiry()
            except Exception as e:
                logger.error(f"Garantie daemon error: {e}")
            try:
                check_contrat_expiry()
            except Exception as e:
                logger.error(f"Contrat daemon error: {e}")
            time.sleep(86400)
    threading.Thread(target=_run, daemon=True, name="notifications-daemon").start()
    logger.info("⏰ Notifications daemon: démarré (garanties + contrats, 24h)")


def check_contrat_expiry():
    """Vérifie les contrats actifs expirant dans les 30 prochains jours."""
    from datetime import date, timedelta
    try:
        df = lire_contrats()
        if df is None or df.empty:
            return []
        today = date.today()
        alert_limit = today + timedelta(days=30)
        alerts = []
        for _, row in df.iterrows():
            statut = str(row.get('statut', '') or '').strip().lower()
            if statut not in ('actif', 'active', ''):
                continue
            date_fin_str = str(row.get('date_fin', '') or '').strip()
            if not date_fin_str:
                continue
            try:
                date_fin = date.fromisoformat(date_fin_str[:10])
                if today <= date_fin <= alert_limit:
                    jours = (date_fin - today).days
                    alerts.append({
                        'id': row.get('id', '?'),
                        'client': row.get('client') or row.get('Client', '?'),
                        'equipement': row.get('equipement', ''),
                        'type_contrat': row.get('type_contrat', ''),
                        'fin': date_fin.strftime('%d/%m/%Y'),
                        'jours': jours,
                    })
            except Exception:
                continue
        if alerts:
            lines = '\n'.join(
                f"  • <b>#{a['id']}</b> {a['client']}"
                + (f" ({a['equipement']})" if a['equipement'] else "")
                + f" — <i>{a['type_contrat']}</i>"
                + f" — expire le {a['fin']} ({a['jours']}j)"
                for a in alerts
            )
            msg = (
                f"📄 <b>Contrats expirant bientôt</b>\n"
                f"{lines}\n"
                f"📅 Vérification SAVIA — {today.strftime('%d/%m/%Y')}"
            )
            _send_telegram(msg)
            logger.info(f"Contrat check: {len(alerts)} alerte(s) envoyée(s)")
        return alerts
    except Exception as e:
        logger.error(f"Contrat expiry check error: {e}")
        return []


def _df_to_records(df) -> list:
    """Convert DataFrame to JSON-safe list of dicts."""
    if df is None or df.empty:
        return []
    records = df.to_dict(orient="records")
    for r in records:
        for k, v in r.items():
            if isinstance(v, float) and math.isnan(v):
                r[k] = None
            elif hasattr(v, 'isoformat'):
                r[k] = v.isoformat()
    return records


def _send_telegram(message: str) -> bool:
    """
    Envoie un message Telegram en lisant token + chat_id depuis la DB config.
    Utilise urllib (pas de dépendance externe).
    """
    import urllib.request, urllib.parse, json as _json
    try:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT cle, valeur FROM config_client WHERE cle = ANY(%s)",
                (["telegram_token", "telegram_chat_id"],)
            ).fetchall()
        config = {r["cle"]: r["valeur"] for r in rows}
        token   = config.get("telegram_token", "").strip()
        chat_id = config.get("telegram_chat_id", "").strip()
        if not token or not chat_id:
            logger.warning("Telegram non configuré — notification ignorée")
            return False

        url  = f"https://api.telegram.org/bot{token}/sendMessage"
        data = urllib.parse.urlencode({
            "chat_id":    chat_id,
            "text":       message,
            "parse_mode": "HTML",
        }).encode("utf-8")
        req = urllib.request.Request(url, data=data)
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = _json.loads(resp.read().decode())
        if result.get("ok"):
            logger.info("Message Telegram envoyé")
            return True
        logger.error(f"Telegram API error: {result}")
        return False
    except Exception as e:
        logger.error(f"Erreur Telegram: {e}")
        return False


def _verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """Verify JWT and return user payload. Returns guest user if no token (allows public read)."""
    if not credentials:
        return {"sub": "guest", "role": "Lecteur", "nom": "Visiteur"}
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        # Rétrocompatibilité : si Lecteur mais client absent du token, le récupérer en DB
        if payload.get("role") == "Lecteur" and not payload.get("client"):
            try:
                with get_db() as conn:
                    row = conn.execute(
                        "SELECT client FROM utilisateurs WHERE username = %s",
                        (payload.get("sub", ""),)
                    ).fetchone()
                    if row and row["client"]:
                        payload["client"] = row["client"]
            except Exception:
                pass
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expiré")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Token invalide")



def _optional_token(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[dict]:
    """Verify JWT if present, return None if missing (allows public read access)."""
    if not credentials:
        return None
    try:
        return jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
    except Exception:
        return None


def _verify_password(password: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


# ==========================================
# AUTH
# ==========================================

class LoginRequest(BaseModel):
    username: str
    password: str


@app.get("/")
def root():
    return {"status": "ok", "service": "SAVIA API", "version": "2.0.0"}


@app.post("/api/auth/login")
def login(body: LoginRequest):
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM utilisateurs WHERE username = %s AND actif = 1",
            (body.username,)
        ).fetchone()

    if not row or not _verify_password(body.password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="Identifiants incorrects")

    user_data = dict(row)
    payload = {
        "sub": user_data["username"],
        "role": user_data["role"],
        "nom": user_data.get("nom_complet", ""),
        "client": user_data.get("client", "") or "",  # ← ajouté
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRY_HOURS),
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")

    return {
        "token": token,
        "user": {
            "username": user_data["username"],
            "nom": user_data.get("nom_complet", ""),
            "role": user_data["role"],
            "client": user_data.get("client", "") or "",  # ← ajouté
        }
    }


@app.get("/api/auth/me")
def me(user: dict = Depends(_verify_token)):
    return {"user": user}


def _get_client_filter(user: dict) -> Optional[str]:
    """Retourne le client restricté pour un Lecteur, None sinon (accès total)."""
    if user.get("role") == "Lecteur":
        c = user.get("client", "").strip()
        return c if c else None
    return None


# ==========================================
# DASHBOARD — Aggregated KPIs
# ==========================================

@app.get("/api/dashboard/kpis")
def get_dashboard_kpis(
    client: Optional[str] = None,
    date_start: Optional[str] = None,
    date_end: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    """Compute real KPIs from the database, optionally filtered by client and date range."""
    try:
        df_eq = lire_equipements()
        df_int = lire_interventions()

        # Pour Lecteur : forcer le filtre par son client
        effective_client = _get_client_filter(user) or client

        # Filter equipements by client
        if effective_client and not df_eq.empty and "Client" in df_eq.columns:
            df_eq = df_eq[df_eq["Client"].astype(str).str.lower() == effective_client.lower()]

        # Filter interventions by client (via matching machines)
        if effective_client and not df_eq.empty and not df_int.empty and "machine" in df_int.columns:
            machines_client = df_eq["Nom"].tolist() if "Nom" in df_eq.columns else []
            df_int = df_int[df_int["machine"].isin(machines_client)]

        # Filter interventions by date range
        if not df_int.empty and "date" in df_int.columns:
            df_int["date"] = pd.to_datetime(df_int["date"], errors="coerce")
            if date_start:
                df_int = df_int[df_int["date"] >= pd.to_datetime(date_start)]
            if date_end:
                df_int = df_int[df_int["date"] <= pd.to_datetime(date_end)]

        nb_eq = len(df_eq) if not df_eq.empty else 0
        nb_critiques = 0
        if not df_eq.empty and "Statut" in df_eq.columns:
            nb_critiques = len(df_eq[df_eq["Statut"].isin(["Hors Service", "Critique"])])

        nb_interventions = len(df_int) if not df_int.empty else 0
        cout_total = 0.0
        mttr = 0.0
        if not df_int.empty:
            # Exclure Installation et Formation des KPIs de maintenance
            TRACABILITE = ['installation', 'formation']
            df_maint = df_int
            if "type_intervention" in df_int.columns:
                df_maint = df_int[~df_int["type_intervention"].str.lower().isin(TRACABILITE)]
            if "cout" in df_maint.columns:
                cout_total = float(df_maint["cout"].sum()) if df_maint["cout"].notna().any() else 0
            if "duree_minutes" in df_maint.columns:
                durees = df_maint["duree_minutes"].dropna()
                mttr = round(float(durees.mean()) / 60, 1) if len(durees) > 0 else 0

        # Disponibilité = % équipements opérationnels
        dispo = 100.0
        if not df_eq.empty and "Statut" in df_eq.columns:
            op = len(df_eq[~df_eq["Statut"].isin(["Hors Service", "Critique"])])
            dispo = round((op / nb_eq) * 100, 1) if nb_eq > 0 else 100

        # MTBF approximation
        mtbf = 720  # default
        if nb_interventions > 0 and nb_eq > 0:
            mtbf = round((nb_eq * 30 * 24) / max(nb_interventions, 1))

        return {
            "nb_equipements": nb_eq,
            "nb_critiques": nb_critiques,
            "disponibilite": dispo,
            "mtbf": mtbf,
            "mttr": mttr,
            "cout_total": round(cout_total, 2),
            "nb_interventions": nb_interventions,
        }
    except Exception as e:
        logger.error(f"Dashboard KPIs error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/dashboard/health-scores")
def get_health_scores(
    client: Optional[str] = None,
    date_start: Optional[str] = None,
    date_end: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    """Compute health scores per equipment based on intervention history."""
    try:
        df_eq = lire_equipements()
        df_int = lire_interventions()

        # Pour Lecteur : forcer le filtre par son client
        effective_client = _get_client_filter(user) or client

        # Filter equipements by client
        if effective_client and not df_eq.empty and "Client" in df_eq.columns:
            df_eq = df_eq[df_eq["Client"].astype(str).str.lower() == effective_client.lower()]

        # Filter interventions by date range
        if not df_int.empty and "date" in df_int.columns:
            df_int["date"] = pd.to_datetime(df_int["date"], errors="coerce")
            if date_start:
                df_int = df_int[df_int["date"] >= pd.to_datetime(date_start)]
            if date_end:
                df_int = df_int[df_int["date"] <= pd.to_datetime(date_end)]

        scores = []

        if df_eq.empty:
            return []

        # Compute period duration in months (for rate-based scoring)
        period_months = 12.0  # default: 1 year
        if date_start and date_end:
            import datetime as _dt
            try:
                d0 = _dt.date.fromisoformat(str(date_start))
                d1 = _dt.date.fromisoformat(str(date_end))
                days = max(1, (d1 - d0).days + 1)
                period_months = max(0.1, days / 30.0)
            except Exception:
                period_months = 12.0

        # Exclude tracabilite interventions from panne count
        TRACABILITE = {"installation", "formation"}
        df_sav = df_int.copy()
        if not df_sav.empty and "type_intervention" in df_sav.columns:
            df_sav = df_sav[~df_sav["type_intervention"].str.lower().isin(TRACABILITE)]

        seen_keys = set()
        for _, eq in df_eq.iterrows():
            nom = eq.get("Nom", "")
            client_val = str(eq.get("Client", "") or "")
            dedup_key = (nom.lower(), client_val.lower())
            if dedup_key in seen_keys:
                continue
            seen_keys.add(dedup_key)

            pannes = 0
            if not df_sav.empty and "machine" in df_sav.columns:
                pannes = len(df_sav[df_sav["machine"] == nom])

            # Rate-based scoring: pannes per month
            panne_rate = pannes / period_months if period_months > 0 else pannes
            if panne_rate <= 0:
                score = 100
            elif panne_rate <= 0.25:
                score = 90
            elif panne_rate <= 0.5:
                score = 78
            elif panne_rate <= 1.0:
                score = 65
            elif panne_rate <= 2.0:
                score = 48
            elif panne_rate <= 4.0:
                score = 30
            elif panne_rate <= 8.0:
                score = 18
            else:
                score = 10

            tendance = "stable"
            if panne_rate > 1.0:
                tendance = "baisse"
            elif panne_rate == 0:
                tendance = "hausse"

            scores.append({
                "machine": nom,
                "score": score,
                "tendance": tendance,
                "pannes": pannes,
                "client": client_val,
            })

        return sorted(scores, key=lambda x: x["score"])
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# ÉQUIPEMENTS
# ==========================================

@app.get("/api/equipements")
def get_equipements(client: Optional[str] = None, user: dict = Depends(_verify_token)):
    df = lire_equipements()
    # Priority: explicit ?client= param, then user's client filter
    client_filter = client or _get_client_filter(user)
    if client_filter and not df.empty and "Client" in df.columns:
        df = df[df["Client"].astype(str).str.lower() == client_filter.lower()]
    return _df_to_records(df)


@app.post("/api/equipements")
def create_equipement(body: dict, user: dict = Depends(_verify_token)):
    ajouter_equipement(body)
    # Return the ID of the created/upserted equipment
    nom = body.get("Nom", "")
    client = body.get("Client", "Centre Principal")
    with get_db() as conn:
        row = conn.execute(
            "SELECT id FROM equipements WHERE nom = ? AND client = ?",
            (nom, client)
        ).fetchone()
    equip_id = dict(row)["id"] if row else None
    return {"ok": True, "id": equip_id}


@app.put("/api/equipements/{equip_id}")
def update_equipement(equip_id: int, body: dict, user: dict = Depends(_verify_token)):
    modifier_equipement(equip_id, body)
    return {"ok": True}


@app.delete("/api/equipements/{equip_id}")
def delete_equipement(equip_id: int, user: dict = Depends(_verify_token)):
    supprimer_equipement(equip_id)
    return {"ok": True}


# ==========================================
# DOCUMENTS TECHNIQUES
# ==========================================

@app.post("/api/documents-techniques/upload")
def upload_document(body: dict, user: dict = Depends(_verify_token)):
    """Upload a technical document (base64 encoded) for an equipment."""
    from db_engine import ajouter_document_technique
    equip_id = body.get("equipement_id")
    nom_fichier = body.get("nom_fichier", "")
    contenu_base64 = body.get("contenu_base64", "")
    if not equip_id or not nom_fichier or not contenu_base64:
        raise HTTPException(status_code=400, detail="equipement_id, nom_fichier et contenu_base64 requis")
    ajouter_document_technique(equip_id, nom_fichier, contenu_base64)
    return {"ok": True}


@app.get("/api/documents-techniques")
def get_all_documents(user: dict = Depends(_verify_token)):
    """List all technical documents with associated equipment info."""
    from db_engine import lire_tous_documents_techniques
    return lire_tous_documents_techniques()


@app.get("/api/documents-techniques/{equip_id}")
def get_documents_by_equipment(equip_id: int, user: dict = Depends(_verify_token)):
    """List technical documents for a specific equipment."""
    from db_engine import lire_documents_techniques
    return lire_documents_techniques(equip_id)


@app.get("/api/documents-techniques/download/{doc_id}")
def download_document(doc_id: int, user: dict = Depends(_verify_token)):
    """Download a specific technical document (returns base64 content)."""
    from db_engine import lire_document_technique_contenu
    doc = lire_document_technique_contenu(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    return doc


@app.delete("/api/documents-techniques/{doc_id}")
def delete_document(doc_id: int, user: dict = Depends(_verify_token)):
    """Delete a technical document."""
    from db_engine import supprimer_document_technique
    supprimer_document_technique(doc_id)
    return {"ok": True}


# ==========================================
# INTERVENTIONS / SAV
# ==========================================

@app.get("/api/interventions")
def get_interventions(
    machine: Optional[str] = None,
    technicien: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    df = lire_interventions(machine=machine)
    # Si le user est un Technicien → filtrer automatiquement ses interventions
    if user.get("role") == "Technicien" and not df.empty:
        user_nom_complet = (user.get("nom") or "").strip()
        # Découper en mots individuels → cherche TOUS les mots dans le champ technicien
        # Gère "Dridi Ali" vs "Ali Dridi" et autres variations d'ordre
        name_words = [w.lower() for w in user_nom_complet.split() if len(w) > 1]
        if name_words and "technicien" in df.columns:
            df = df[df["technicien"].astype(str).apply(
                lambda t: all(word in t.lower() for word in name_words)
            )]
        elif not name_words:
            # Aucun nom disponible → ne rien filtrer (afficher tout)
            pass
    elif technicien and not df.empty and "technicien" in df.columns:
        words = technicien.lower().split()
        df = df[df["technicien"].astype(str).apply(
            lambda t: all(w in t.lower() for w in words)
        )]
    # Filtrage par client pour Lecteur
    client_filter = _get_client_filter(user)
    if client_filter and not df.empty:
        df_eq = lire_equipements()
        if not df_eq.empty and "Client" in df_eq.columns and "Nom" in df_eq.columns:
            machines_client = set(
                df_eq[df_eq["Client"].astype(str).str.lower() == client_filter.lower()]["Nom"].tolist()
            )
            if "machine" in df.columns:
                df = df[df["machine"].isin(machines_client)]
    return _df_to_records(df)


@app.post("/api/interventions")
def create_intervention(body: dict, user: dict = Depends(_verify_token)):
    ajouter_intervention(body)
    return {"ok": True}


@app.put("/api/interventions/{intervention_id}")
def update_intervention(intervention_id: int, body: dict, user: dict = Depends(_verify_token)):
    new_statut = body.get("statut")
    if new_statut and "tur" in new_statut.lower():
        # Normaliser pieces_a_deduire : s'assurer que c'est une liste de dicts avec clé 'ref'
        raw_pieces = body.get("pieces_a_deduire") or []
        if not isinstance(raw_pieces, list):
            raw_pieces = []
        pieces_valides = [p for p in raw_pieces if isinstance(p, dict) and p.get("ref")]

        try:
            ok, msg = cloturer_intervention(
                intervention_id,
                body.get("probleme", ""),
                body.get("cause", ""),
                body.get("solution", ""),
                pieces_a_deduire=pieces_valides if pieces_valides else None,
                duree_minutes=body.get("duree_minutes", 0),
            )
            if not ok:
                raise HTTPException(status_code=400, detail=msg)

            # --- Telegram notification clôture ---
            try:
                with get_db() as conn:
                    row = conn.execute(
                        "SELECT machine, technicien, probleme, cause, solution, duree_minutes, notes FROM interventions WHERE id = %s",
                        (intervention_id,)
                    ).fetchone()
                if row:
                    d = dict(row)
                    duree_h = round((d.get('duree_minutes') or 0) / 60, 1)
                    notes_line = f"\n📌 Notes : {d['notes']}" if d.get('notes') else ""
                    msg_tg = (
                        f"✅ <b>INTERVENTION CLÔTURÉE — #{intervention_id}</b>\n\n"
                        f"🏥 Machine : <b>{d.get('machine', '')}</b>\n"
                        f"👷 Technicien : <b>{d.get('technicien', '')}</b>\n"
                        f"🔴 Problème : {str(d.get('probleme', ''))[:200]}\n"
                        f"🔍 Cause : {str(d.get('cause', ''))[:200]}\n"
                        f"🟢 Solution : {str(d.get('solution', ''))[:200]}\n"
                        f"⏱️ Durée : <b>{duree_h}h</b>"
                        f"{notes_line}\n"
                        f"🕐 {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                    )
                    _send_telegram(msg_tg)
            except Exception as te:
                logger.error(f"Telegram clôture erreur: {te}")

            # --- Mettre à jour la demande liée (si elle existe) → statut "Résolue" ---
            try:
                with get_db() as conn:
                    conn.execute(
                        """UPDATE demandes_intervention
                           SET statut = 'Résolue',
                               date_traitement = %s
                         WHERE intervention_id = %s
                           AND statut != 'Résolue'""",
                        (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), intervention_id)
                    )
                    logger.info(f"Demande liée à l'intervention #{intervention_id} marquée Résolue")
            except Exception as de:
                logger.error(f"Erreur mise à jour demande liée: {de}")

            return {"ok": True, "message": msg}
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Erreur clôture intervention #{intervention_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Erreur lors de la clôture: {str(e)}")
    if new_statut and "attente" in new_statut.lower() and "pi" in new_statut.lower():
        # Statut = "En attente de pièce" → notification rupture pour gestionnaires
        pieces_attente = body.get("pieces_rupture") or []
        try:
            with get_db() as conn:
                row = conn.execute(
                    "SELECT machine, technicien, notes FROM interventions WHERE id = %s",
                    (intervention_id,)
                ).fetchone()
            if row:
                machine = row["machine"] or ""
                technicien = row["technicien"] or ""
                # Extraire client depuis notes [Client]
                notes = str(row.get("notes") or "")
                client = notes[1:notes.index("]")] if notes.startswith("[") and "]" in notes else ""
                for piece in pieces_attente:
                    ref = piece.get("reference") or piece.get("ref") or ""
                    nom = piece.get("designation") or piece.get("nom") or ref
                    ajouter_notification_piece({
                        "type": "piece_rupture",
                        "intervention_id": intervention_id,
                        "piece_reference": ref,
                        "piece_nom": nom,
                        "intervention_ref": f"#{intervention_id}",
                        "equipement": machine,
                        "client": client,
                        "technicien": technicien,
                        "message": f"⚠️ Intervention #{intervention_id} sur {machine} en attente de la pièce "
                                   f"{ref} ({nom}) — rupture de stock",
                        "source": "sav",
                        "destination": "gestionnaire",
                    })
                    logger.info(f"Notif rupture créée: pièce {ref} pour intervention #{intervention_id}")
        except Exception as ne:
            logger.error(f"Erreur création notif rupture: {ne}")

    if new_statut:
        update_intervention_statut(intervention_id, new_statut)
    # Update other fields
    fields = []
    params = []
    for f in ["probleme", "cause", "solution", "pieces_utilisees", "cout",
              "duree_minutes", "description", "notes", "type_erreur", "priorite",
              "fiche_validation"]:
        if f in body:
            fields.append(f"{f} = %s")
            params.append(body[f])
    if fields:
        params.append(intervention_id)
        with get_db() as conn:
            conn.execute(f"UPDATE interventions SET {', '.join(fields)} WHERE id = %s", params)
    return {"ok": True}



@app.post("/api/interventions/{intervention_id}/fiche")
async def upload_fiche(intervention_id: int, file: UploadFile = File(...), user: dict = Depends(_verify_token)):
    """Upload la photo de la fiche signée pour une intervention clôturée."""
    contents = await file.read()
    with get_db() as conn:
        conn.execute(
            "UPDATE interventions SET fiche_photo_nom = %s, fiche_photo_data = %s WHERE id = %s",
            (file.filename, contents, intervention_id)
        )
    logger.info(f"Fiche photo uploadée pour intervention #{intervention_id}: {file.filename}")
    return {"ok": True, "filename": file.filename}


@app.post("/api/interventions/{intervention_id}/photo")
async def upload_photo_alias(intervention_id: int,
                             photo: UploadFile = File(None),
                             file: UploadFile = File(None),
                             user: dict = Depends(_verify_token)):
    """Alias /photo → /fiche pour compatibilité avec l'ancien api_server.py (Streamlit).
    Accepte le champ 'photo' ou 'file'."""
    upload = photo or file
    if not upload:
        raise HTTPException(status_code=400, detail="Aucun fichier fourni")
    contents = await upload.read()
    with get_db() as conn:
        conn.execute(
            "UPDATE interventions SET fiche_photo_nom = %s, fiche_photo_data = %s WHERE id = %s",
            (upload.filename, contents, intervention_id)
        )
    logger.info(f"[/photo alias] Fiche photo uploadée pour intervention #{intervention_id}: {upload.filename}")
    return {"ok": True, "message": "Photo enregistrée", "filename": upload.filename}


@app.get("/api/interventions/{intervention_id}/fiche")
def download_fiche(intervention_id: int, token: Optional[str] = Query(None), user: dict = Depends(_verify_token)):
    """Télécharge la photo de fiche pour une intervention (accepte token en query param pour img src)."""
    from fastapi.responses import Response
    # Si token en query param, valider manuellement
    if token:
        try:
            jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        except Exception:
            raise HTTPException(status_code=401, detail="Token invalide")
    with get_db() as conn:
        try:
            row = conn.execute(
                "SELECT fiche_photo_nom, fiche_photo_data FROM interventions WHERE id = %s",
                (intervention_id,)
            ).fetchone()
        except Exception:
            raise HTTPException(status_code=404, detail="Colonne fiche non trouvée")
    if not row or not row["fiche_photo_data"]:
        raise HTTPException(status_code=404, detail="Aucune fiche pour cette intervention")
    nom = row["fiche_photo_nom"] or f"fiche_{intervention_id}.jpg"
    data = bytes(row["fiche_photo_data"])
    ext = nom.rsplit('.', 1)[-1].lower() if '.' in nom else 'jpg'
    mime_map = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
                'pdf': 'application/pdf', 'webp': 'image/webp'}
    mime = mime_map.get(ext, 'application/octet-stream')
    return Response(content=data, media_type=mime,
                    headers={"Content-Disposition": f'inline; filename="{nom}"'})



@app.get("/api/interventions/fiches")
def list_fiches(user: dict = Depends(_verify_token)):
    """Liste uniquement les interventions clôturées AVEC photo attachée."""
    with get_db() as conn:
        try:
            rows = conn.execute("""
                SELECT id, date, machine, technicien, statut, probleme, solution,
                       duree_minutes,
                       COALESCE(fiche_photo_nom, '') AS fiche_photo_nom,
                       (fiche_photo_data IS NOT NULL AND octet_length(fiche_photo_data) > 0) AS has_fiche,
                       COALESCE(fiche_validation, 'En attente') AS fiche_validation
                FROM interventions
                WHERE (statut ILIKE '%lotur%' OR statut ILIKE '%termin%' OR statut = 'Cloturee')
                  AND fiche_photo_data IS NOT NULL
                  AND octet_length(fiche_photo_data) > 0
                ORDER BY id DESC
                LIMIT 200
            """).fetchall()
        except Exception as e:
            logger.error(f"Erreur list_fiches: {e}")
            return []
    result = []
    for r in rows:
        d = dict(r)
        if hasattr(d.get('date'), 'isoformat'):
            d['date'] = d['date'].isoformat()
        d['has_fiche'] = bool(d.get('has_fiche'))
        d['fiche_validation'] = d.get('fiche_validation') or 'En attente'
        result.append(d)
    return result


@app.patch("/api/interventions/{intervention_id}/fiche-validation")
def update_fiche_validation(intervention_id: int, body: dict, user: dict = Depends(_verify_token)):
    """Met à jour le statut de validation client d'une fiche.
    Une fois 'Validée', aucune modification n'est plus possible."""
    nouveau_statut = body.get("validation", "").strip()
    valeurs_autorisees = {"En attente", "Validée"}
    if nouveau_statut not in valeurs_autorisees:
        raise HTTPException(status_code=400, detail=f"Valeur invalide: {nouveau_statut}. Valeurs autorisées: {valeurs_autorisees}")

    with get_db() as conn:
        # Vérifier le statut actuel
        row = conn.execute(
            "SELECT fiche_validation FROM interventions WHERE id = %s",
            (intervention_id,)
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Intervention non trouvée")
        statut_actuel = (row["fiche_validation"] or "En attente").strip()
        if statut_actuel == "Validée":
            raise HTTPException(status_code=403, detail="Fiche déjà validée — aucune modification possible")
        conn.execute(
            "UPDATE interventions SET fiche_validation = %s WHERE id = %s",
            (nouveau_statut, intervention_id)
        )
    logger.info(f"Fiche #{intervention_id}: validation mise à jour → '{nouveau_statut}' par {user.get('nom', '?')}")
    return {"ok": True, "validation": nouveau_statut}


# ==========================================
# DEMANDES D'INTERVENTION
# ==========================================

@app.get("/api/demandes")
def get_demandes(
    statuts: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    df = lire_demandes_intervention()
    # Lecteur : ne voit que les demandes de son client
    client_filter = _get_client_filter(user)
    if client_filter and not df.empty and "client" in df.columns:
        df = df[df["client"].astype(str).str.lower() == client_filter.lower()]
    if statuts and not df.empty and "statut" in df.columns:
        lst = [s.strip() for s in statuts.split(",")]
        df = df[df["statut"].isin(lst)]
    return _df_to_records(df)


@app.post("/api/demandes")
def create_demande(body: dict, user: dict = Depends(_verify_token)):
    from db_engine import get_db
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    demandeur          = body.get("demandeur") or user.get("username", "")
    client             = body.get("client") or ""
    equipement         = body.get("equipement") or ""
    urgence            = body.get("urgence") or "Moyenne"
    description        = body.get("description") or ""
    code_erreur        = body.get("code_erreur") or ""
    contact_nom        = body.get("contact_nom") or ""
    contact_tel        = body.get("contact_tel") or ""
    technicien_assigne = body.get("technicien_assigne") or ""
    # Si technicien assigné dès la création → statut "Assignée"
    statut = body.get("statut") or ("Assignée" if technicien_assigne else "En attente")

    with get_db() as conn:
        conn.execute("""
            INSERT INTO demandes_intervention
              (date_demande, demandeur, client, equipement, urgence,
               description, code_erreur, contact_nom, contact_tel,
               statut, technicien_assigne)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            body.get("date_demande") or now_str,
            demandeur, client, equipement, urgence,
            description, code_erreur, contact_nom, contact_tel,
            statut, technicien_assigne,
        ))
        # Récupérer l'id de la demande créée
        new_demande = conn.execute(
            "SELECT id FROM demandes_intervention ORDER BY id DESC LIMIT 1"
        ).fetchone()
        demande_id = new_demande["id"] if new_demande else None

    # --- Notification Telegram ---
    urg_icon = "\U0001f534" if urgence in ("Haute", "Critique") else "\U0001f7e1" if urgence == "Moyenne" else "\U0001f7e2"
    contact_line = f"\n\U0001f4de Contact : <b>{contact_nom}</b>" + (f" — {contact_tel}" if contact_tel else "") if contact_nom else ""
    code_line    = f"\n\U0001f522 Code erreur : <code>{code_erreur}</code>" if code_erreur else ""
    tech_line    = f"\n\U0001f477 Assigné à : <b>{technicien_assigne}</b>" if technicien_assigne else ""
    msg = (
        f"\U0001f4cb <b>NOUVELLE DEMANDE D'INTERVENTION</b>\n\n"
        f"\U0001f3e2 Client : <b>{client}</b>\n"
        f"\U0001f3e5 Équipement : <b>{equipement}</b>\n"
        f"{urg_icon} Urgence : <b>{urgence}</b>\n"
        f"\U0001f4dd Problème : {description[:300]}"
        f"{code_line}"
        f"{contact_line}"
        f"{tech_line}\n"
        f"\U0001f464 Demandeur : <b>{demandeur}</b>\n"
        f"\U0001f550 Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
        f"\U0001f449 Connectez-vous à <b>SAVIA</b> pour traiter cette demande."
    )
    _send_telegram(msg)

    # --- Auto-créer une intervention SAV si technicien assigné dès la création ---
    if technicien_assigne and demande_id:
        try:
            with get_db() as conn:
                today = datetime.now().strftime("%Y-%m-%d")
                notes_interv = f"[{client}] Demande #{demande_id}"
                conn.execute("""
                    INSERT INTO interventions
                      (date, machine, technicien, type_intervention, description,
                       probleme, code_erreur, statut, priorite, notes)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    today,
                    equipement,
                    technicien_assigne,
                    "Corrective",
                    description[:500],
                    description[:500],
                    code_erreur,
                    "En cours",
                    urgence,
                    notes_interv,
                ))
                new_interv = conn.execute(
                    "SELECT id FROM interventions ORDER BY id DESC LIMIT 1"
                ).fetchone()
                if new_interv:
                    conn.execute(
                        "UPDATE demandes_intervention SET intervention_id = %s WHERE id = %s",
                        (new_interv["id"], demande_id)
                    )
                    logger.info(f"Intervention #{new_interv['id']} auto-créée pour demande #{demande_id} → {technicien_assigne}")
        except Exception as e:
            logger.error(f"Erreur auto-création intervention depuis demande: {e}")

    return {"success": True}


@app.put("/api/demandes/{demande_id}/statut")
def update_demande_statut(demande_id: int, body: dict, user: dict = Depends(_verify_token)):
    from db_engine import get_db
    nouveau_statut      = body.get("statut") or "En cours"
    technicien_assigne  = body.get("technicien_assigne") or ""
    notes_traitement    = body.get("notes_traitement") or ""

    # Récupérer les données de la demande AVANT mise à jour pour le message
    demande_info = {}
    with get_db() as conn:
        row = conn.execute(
            "SELECT client, equipement, urgence, description, demandeur, contact_nom, contact_tel FROM demandes_intervention WHERE id = %s",
            (demande_id,)
        ).fetchone()
        if row:
            demande_info = dict(row)
        conn.execute("""
            UPDATE demandes_intervention
            SET statut = %s, technicien_assigne = %s, notes_traitement = %s,
                date_traitement = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (nouveau_statut, technicien_assigne, notes_traitement, demande_id))

    # --- Notification Telegram (tous les changements de statut) ---
    statut_icons = {
        "En attente": "\u23f3",
        "En cours":   "\U0001f527",
        "Assign\u00e9e": "\U0001f477",
        "R\u00e9solue":  "\u2705",
        "Cl\u00f4tur\u00e9e": "\U0001f3c1",
        "Plan\u00e0fi\u00e9e": "\U0001f4c5",
        "Planifi\u00e9e":  "\U0001f4c5",
        "Rejet\u00e9e":    "\u274c",
        "Accept\u00e9e":  "\u2705",
    }
    icon = statut_icons.get(nouveau_statut, "\U0001f4cb")
    client     = demande_info.get("client", "")
    equipement = demande_info.get("equipement", "")
    urgence    = demande_info.get("urgence", "")
    description = str(demande_info.get("description", ""))[:300]
    demandeur  = demande_info.get("demandeur", "")
    contact_nom = demande_info.get("contact_nom", "")
    contact_tel = demande_info.get("contact_tel", "")

    urg_icon     = "\U0001f534" if urgence in ("Haute", "Critique") else "\U0001f7e1" if urgence == "Moyenne" else "\U0001f7e2"
    tech_line    = f"\n\U0001f477 Technicien : <b>{technicien_assigne}</b>" if technicien_assigne else ""
    notes_line   = f"\n\U0001f4cc Notes : {notes_traitement}" if notes_traitement else ""
    contact_line = f"\n\U0001f4de Contact : <b>{contact_nom}</b>" + (f" — {contact_tel}" if contact_tel else "") if contact_nom else ""
    demandeur_line = f"\n\U0001f464 Demandeur : <b>{demandeur}</b>" if demandeur else ""

    msg = (
        f"{icon} <b>DEMANDE #{demande_id} \u2014 MISE \u00c0 JOUR STATUT</b>\n\n"
        f"\U0001f3e2 Client : <b>{client}</b>\n"
        f"\U0001f3e5 \u00c9quipement : <b>{equipement}</b>\n"
        f"{urg_icon} Urgence : <b>{urgence}</b>\n"
        f"\U0001f4ca Statut : <b>{nouveau_statut}</b>\n"
        f"\U0001f4dd Probl\u00e8me : {description}"
        f"{tech_line}"
        f"{notes_line}"
        f"{contact_line}"
        f"{demandeur_line}\n"
        f"\U0001f550 {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    _send_telegram(msg)

    # --- Auto-créer une intervention si technicien assigné et pas déjà liée ---
    if technicien_assigne:
        try:
            with get_db() as conn:
                # Vérifier si une intervention existe déjà pour cette demande
                existing = conn.execute(
                    "SELECT intervention_id FROM demandes_intervention WHERE id = %s",
                    (demande_id,)
                ).fetchone()
                already_linked = existing and existing["intervention_id"]

                if not already_linked:
                    # Créer l'intervention
                    today = datetime.now().strftime("%Y-%m-%d")
                    notes_interv = f"[{client}] Demande #{demande_id}"
                    if notes_traitement:
                        notes_interv += f" — {notes_traitement}"
                    conn.execute("""
                        INSERT INTO interventions
                          (date, machine, technicien, type_intervention, description,
                           probleme, code_erreur, statut, priorite, notes)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        today,
                        equipement,
                        technicien_assigne,
                        "Corrective",
                        description[:500],
                        description[:500],
                        demande_info.get("code_erreur", "") or "",
                        "En cours",
                        urgence,
                        notes_interv,
                    ))
                    # Récupérer l'id de l'intervention créée
                    new_interv = conn.execute(
                        "SELECT id FROM interventions ORDER BY id DESC LIMIT 1"
                    ).fetchone()
                    if new_interv:
                        conn.execute(
                            "UPDATE demandes_intervention SET intervention_id = %s WHERE id = %s",
                            (new_interv["id"], demande_id)
                        )
                        logger.info(f"Intervention #{new_interv['id']} auto-créée pour demande #{demande_id} → {technicien_assigne}")
        except Exception as e:
            logger.error(f"Erreur auto-création intervention: {e}")

    return {"success": True}


# ==========================================
# PIÈCES DE RECHANGE
# ==========================================

@app.get("/api/pieces")
def get_pieces(user: dict = Depends(_verify_token)):
    return _df_to_records(lire_pieces())


@app.post("/api/pieces")
def create_piece(body: dict, user: dict = Depends(_verify_token)):
    ajouter_piece(body)
    return {"ok": True}


@app.put("/api/pieces/{piece_id}")
def update_piece(piece_id: int, body: dict, user: dict = Depends(_verify_token)):
    """Mise à jour d'une pièce. Si stock passe de 0 → >0, déclenche notifications pour les techniciens en attente."""
    # Récupérer le stock AVANT modification pour détecter le réapprovisionnement
    nouveau_stock = body.get("stock_actuel")
    try:
        with get_db() as conn:
            old = conn.execute(
                "SELECT reference, designation, stock_actuel FROM pieces_rechange WHERE id = %s",
                (piece_id,)
            ).fetchone()
        stock_avant = int(old["stock_actuel"]) if old else None
        reference = old["reference"] if old else ""
        nom_piece = old["designation"] if old else ""
    except Exception:
        stock_avant = None
        reference = ""
        nom_piece = ""

    modifier_piece(piece_id, body)

    # Détecter réapprovisionnement : stock passe de 0 (ou négatif) → positif
    if nouveau_stock is not None and stock_avant is not None:
        try:
            if int(stock_avant) <= 0 and int(nouveau_stock) > 0 and reference:
                # Chercher toutes les notifications rupture non traitées pour cette pièce
                df_notifs = notifications_rupture_pour_piece(reference)
                if not df_notifs.empty:
                    # Grouper par technicien
                    tech_map: dict = {}
                    for _, n in df_notifs.iterrows():
                        t = n.get("technicien") or "inconnu"
                        if t not in tech_map:
                            tech_map[t] = []
                        tech_map[t].append({
                            "machine": n.get("equipement") or "",
                            "intervention_id": n.get("intervention_id") or "",
                        })

                    for tech, interventions_list in tech_map.items():
                        machines = ", ".join(set(i["machine"] for i in interventions_list if i["machine"]))
                        nb = len(interventions_list)
                        inter_ids = ", ".join(
                            f"#{i['intervention_id']}" for i in interventions_list if i.get("intervention_id")
                        )
                        ajouter_notification_piece({
                            "type": "piece_dispo",
                            "piece_reference": reference,
                            "piece_nom": nom_piece,
                            "technicien": tech,
                            "equipement": machines,
                            "message": (
                                f"✅ La pièce {reference} ({nom_piece}) est maintenant disponible — "
                                f"{nb} intervention(s) en attente sur : {machines or 'N/A'}"
                            ),
                            "source": "stock",
                            "destination": "technicien",
                        })
                        logger.info(f"Notif piece_dispo créée pour technicien {tech}: pièce {reference}")

                    # --- Telegram : pièce à nouveau disponible ---
                    try:
                        all_techs = ", ".join(tech_map.keys()) or "N/A"
                        all_machines = ", ".join(
                            set(i["machine"] for ivs in tech_map.values() for i in ivs if i.get("machine"))
                        ) or "N/A"
                        all_inter_ids = ", ".join(
                            f"#{i['intervention_id']}"
                            for ivs in tech_map.values() for i in ivs
                            if i.get("intervention_id")
                        ) or "N/A"
                        msg_tg = (
                            f"🟢 <b>PIÈCE DISPONIBLE</b>\n\n"
                            f"🔩 Pièce : <b>{nom_piece}</b>\n"
                            f"🏷 Référence : <b>{reference}</b>\n"
                            f"📦 Stock actuel : <b>{nouveau_stock}</b>\n\n"
                            f"🔗 Intervention(s) concernée(s) : {all_inter_ids}\n"
                            f"🏥 Équipement(s) : {all_machines}\n"
                            f"👷 Technicien(s) : {all_techs}\n"
                            f"🕐 {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                        )
                        _send_telegram(msg_tg)
                        logger.info(f"Telegram pièce disponible envoyé: {reference}")
                    except Exception as tg_err:
                        logger.error(f"Telegram pièce dispo erreur: {tg_err}")

                    # Marquer les notifications rupture comme traitées
                    for _, n in df_notifs.iterrows():
                        try:
                            marquer_notification_traitee(int(n["id"]))
                        except Exception:
                            pass
        except Exception as ne:
            logger.error(f"Erreur notif réappro pièce {piece_id}: {ne}")


    return {"ok": True}


@app.delete("/api/pieces/{piece_id}")
def delete_piece(piece_id: int, user: dict = Depends(_verify_token)):
    supprimer_piece(piece_id)
    return {"ok": True}


# ==========================================
# CONTRATS
# ==========================================

@app.get("/api/contrats")
def get_contrats(client: Optional[str] = None, user: dict = Depends(_verify_token)):
    # Pour Lecteur : forcer le filtre par son client
    effective_client = _get_client_filter(user) or client
    return _df_to_records(lire_contrats(client=effective_client))


@app.post("/api/contrats")
def create_contrat(body: dict, user: dict = Depends(_verify_token)):
    ajouter_contrat(body)
    return {"ok": True}


@app.put("/api/contrats/{contrat_id}")
def update_contrat(contrat_id: int, body: dict, user: dict = Depends(_verify_token)):
    modifier_contrat(contrat_id, body)
    return {"ok": True}


@app.delete("/api/contrats/{contrat_id}")
def delete_contrat(contrat_id: int, user: dict = Depends(_verify_token)):
    supprimer_contrat(contrat_id)
    return {"ok": True}


# ==========================================
# CONFORMITÉ
# ==========================================

@app.get("/api/conformite")
def get_conformite(client: Optional[str] = None, user: dict = Depends(_verify_token)):
    return _df_to_records(lire_conformite(client=client))


@app.post("/api/conformite")
def create_conformite(body: dict, user: dict = Depends(_verify_token)):
    ajouter_conformite(body)
    return {"ok": True}


@app.delete("/api/conformite/{conformite_id}")
def delete_conformite(conformite_id: int, user: dict = Depends(_verify_token)):
    supprimer_conformite(conformite_id)
    return {"ok": True}


# ==========================================
# PLANNING
# ==========================================

@app.get("/api/planning")
def get_planning(
    machine: Optional[str] = None,
    statut: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    df = lire_planning(machine=machine, statut=statut)
    # Pour Lecteur : filtrer par les machines de son client
    client_filter = _get_client_filter(user)
    if client_filter and not df.empty:
        df_eq = lire_equipements()
        if not df_eq.empty and "Client" in df_eq.columns and "Nom" in df_eq.columns:
            machines_client = set(
                df_eq[df_eq["Client"].astype(str).str.lower() == client_filter.lower()]["Nom"].tolist()
            )
            if "machine" in df.columns:
                df = df[df["machine"].isin(machines_client)]
            elif "equipement" in df.columns:
                df = df[df["equipement"].isin(machines_client)]
    return _df_to_records(df)


@app.post("/api/planning")
def create_planning(body: dict, user: dict = Depends(_verify_token)):
    ajouter_planning(body)
    return {"ok": True}


@app.put("/api/planning/{planning_id}")
def update_planning_status(planning_id: int, body: dict, user: dict = Depends(_verify_token)):
    update_planning_statut(planning_id, body.get("statut", ""), body.get("date_realisee"))
    return {"ok": True}


@app.delete("/api/planning/{planning_id}")
def delete_planning(planning_id: int, user: dict = Depends(_verify_token)):
    supprimer_planning(planning_id)
    return {"ok": True}


# ==========================================
# KNOWLEDGE BASE
# ==========================================

@app.get("/api/knowledge")
def get_knowledge(user: dict = Depends(_verify_token)):
    """Return error codes + solutions merged."""
    hex_db, sol_db = lire_base()
    results = []
    for code, info in hex_db.items():
        sol = sol_db.get(code, {})
        results.append({
            "code": code,
            "message": info.get("Msg", ""),
            "level": info.get("Level", ""),
            "type": info.get("Type", ""),
            "cause": sol.get("Cause", ""),
            "solution": sol.get("Solution", ""),
            "priorite": sol.get("Priorité", ""),
        })
    return results


def _parse_text_to_rows(text: str) -> list:
    """Parse unstructured text (from PDF/Word) into error code rows using AI or regex."""
    import re
    rows = []

    # Try AI extraction first
    try:
        from ai_engine import _call_ia, clean_json_response, AI_AVAILABLE
        if AI_AVAILABLE and len(text) > 50:
            prompt = f"""Extrais les codes d'erreur de ce texte technique. 
Pour chaque code trouvé, donne: code, message, type (Hardware/Software/Network), cause, solution, priorite (HAUTE/MOYENNE/BASSE).
Texte (extrait): {text[:4000]}

Réponds en JSON: [{{"code":"ERR001","message":"...","type":"Hardware","cause":"...","solution":"...","priorite":"MOYENNE"}}]"""
            raw = _call_ia(prompt, timeout=30, is_json=True)
            if raw:
                result = clean_json_response(raw)
                if isinstance(result, list) and len(result) > 0:
                    return result
    except Exception:
        pass

    # Fallback: regex-based extraction
    # Common patterns: "ERR-001", "E001", "0x1234", "ERROR 001"
    patterns = [
        r'((?:ERR|ERROR|E|WARN|W|FAULT|F|CODE)[_\-\s]?\d{2,5})',
        r'(0x[0-9A-Fa-f]{4,8})',
        r'((?:H|S|N)\d{4})',
    ]
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            code = match.group(1).strip()
            # Get surrounding context (100 chars)
            start = max(0, match.start() - 20)
            end = min(len(text), match.end() + 200)
            context = text[start:end].replace('\n', ' ').strip()
            if code not in [r.get('code') for r in rows]:
                rows.append({
                    "code": code,
                    "message": context[:120],
                    "type": "Hardware",
                    "cause": "",
                    "solution": "",
                    "priorite": "MOYENNE",
                })

    if not rows:
        # If no codes found, store the document text as a single entry
        lines = [l.strip() for l in text.split('\n') if l.strip() and len(l.strip()) > 10]
        for i, line in enumerate(lines[:50]):
            rows.append({
                "code": f"DOC-{i+1:03d}",
                "message": line[:200],
                "type": "Documentation",
                "cause": "",
                "solution": "",
                "priorite": "BASSE",
            })

    return rows


@app.post("/api/knowledge/import")
async def import_knowledge(file: UploadFile = File(...), user: dict = Depends(_verify_token)):
    """Import error codes from an uploaded Excel/CSV file."""
    import io
    filename = file.filename or ""
    content = await file.read()

    try:
        if filename.endswith(".csv"):
            import csv
            text = content.decode("utf-8", errors="replace")
            reader = csv.DictReader(io.StringIO(text))
            rows = list(reader)
        elif filename.endswith((".xlsx", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
            ws = wb.active
            headers = [str(c.value or "").strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
            rows = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                rows.append({headers[i]: (str(v) if v else "") for i, v in enumerate(row) if i < len(headers)})
        elif filename.endswith(".pdf"):
            # Parse PDF text using PyMuPDF (fitz)
            try:
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                full_text = "\n".join(page.get_text() for page in doc)
            except Exception:
                full_text = content.decode("utf-8", errors="replace")
            rows = _parse_text_to_rows(full_text)
        elif filename.endswith((".docx", ".doc")):
            # Parse Word text
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        full_text += "\n" + " | ".join(cell.text for cell in row.cells)
            except Exception:
                full_text = content.decode("utf-8", errors="replace")
            rows = _parse_text_to_rows(full_text)
        else:
            raise HTTPException(status_code=400, detail="Format non supporté. Utilisez CSV, XLSX, PDF ou DOCX.")

        # Auto-detect column mapping
        col_map = {}
        for h in (rows[0].keys() if rows else []):
            hl = h.lower().strip()
            if "code" in hl: col_map["code"] = h
            elif "message" in hl or "msg" in hl: col_map["message"] = h
            elif "type" in hl: col_map["type"] = h
            elif "cause" in hl: col_map["cause"] = h
            elif "solution" in hl: col_map["solution"] = h
            elif "priorit" in hl: col_map["priorite"] = h

        if "code" not in col_map:
            raise HTTPException(status_code=400, detail="Colonne 'Code' non trouvée dans le fichier.")

        imported = 0
        with get_db() as conn:
            for row in rows:
                code = row.get(col_map.get("code", ""), "").strip()
                if not code:
                    continue
                msg = row.get(col_map.get("message", ""), "")
                typ = row.get(col_map.get("type", ""), "Hardware")
                cause = row.get(col_map.get("cause", ""), "")
                solution = row.get(col_map.get("solution", ""), "")
                priorite = row.get(col_map.get("priorite", ""), "MOYENNE")

                # Insert or update codes_erreurs
                conn.execute(
                    "INSERT INTO codes_erreurs (code, message, type) VALUES (%s, %s, %s) "
                    "ON CONFLICT (code) DO UPDATE SET message=EXCLUDED.message, type=EXCLUDED.type",
                    (code, msg, typ)
                )
                # Insert or update solutions
                conn.execute(
                    "INSERT INTO solutions (code, cause, solution, priorite) VALUES (%s, %s, %s, %s) "
                    "ON CONFLICT (code) DO UPDATE SET cause=EXCLUDED.cause, solution=EXCLUDED.solution, priorite=EXCLUDED.priorite",
                    (code, cause, solution, priorite)
                )
                imported += 1

        return {"ok": True, "imported": imported, "message": f"{imported} codes importés avec succès."}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur d'import: {str(e)}")


# ==========================================
# TECHNICIENS
# ==========================================

@app.get("/api/techniciens")
def get_techniciens(user: dict = Depends(_verify_token)):
    return _df_to_records(lire_techniciens())


@app.post("/api/techniciens")
def create_technicien(body: dict, user: dict = Depends(_verify_token)):
    ajouter_technicien(body)
    return {"ok": True}


@app.put("/api/techniciens/{tech_id}")
def modifier_techniciens_route(tech_id: int, body: dict, user: dict = Depends(_verify_token)):
    update_technicien(tech_id, body)
    return {"ok": True}


@app.delete("/api/techniciens/{tech_id}")
def delete_technicien(tech_id: int, user: dict = Depends(_verify_token)):
    supprimer_technicien(tech_id)
    return {"ok": True}


# ==========================================
# LOGS UPLOAD (Supervision) — S3/MinIO + PostgreSQL
# ==========================================

@app.post("/api/logs/upload")
def upload_log(body: dict, user: dict = Depends(_verify_token)):
    """Upload un fichier log : contenu vers S3/MinIO, métadonnées vers PostgreSQL."""
    import hashlib
    equipement = body.get("equipement", "")
    filename = body.get("filename", "unknown.log")
    content = body.get("content", "")
    nb_errors = body.get("nb_errors", 0)
    nb_critiques = body.get("nb_critiques", 0)
    import json as _json_import
    parsed_errors_raw = body.get("parsed_errors", None)
    parsed_errors_str = _json_import.dumps(parsed_errors_raw) if parsed_errors_raw is not None else None

    if not equipement or not content:
        raise HTTPException(status_code=400, detail="Équipement et contenu requis")

    content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
    username = user.get("sub", "system") if user else "system"

    try:
        with get_db() as conn:
            existing = conn.execute(
                "SELECT id FROM logs_uploaded WHERE content_hash = ? AND equipement = ?",
                (content_hash, equipement)
            ).fetchone()
            if existing:
                eid = existing.get("id") if isinstance(existing, dict) else existing[0]
                # Update parsed_errors on duplicate if not already stored
                if parsed_errors_str:
                    conn.execute(
                        "UPDATE logs_uploaded SET parsed_errors = ? WHERE id = ? AND (parsed_errors IS NULL OR parsed_errors = '')",
                        (parsed_errors_str, eid)
                    )
                return {"ok": True, "message": "Ce log a déjà été enregistré", "id": eid, "duplicate": True}

            # Upload contenu vers S3/MinIO
            s3_key = ""
            size_bytes = len(content.encode('utf-8'))
            try:
                from s3_storage import upload_file as s3_upload
                s3_result = s3_upload(content, filename, equipement, {
                    "nb_errors": str(nb_errors), "uploaded_by": username,
                })
                if s3_result:
                    s3_key = s3_result["s3_key"]
                    size_bytes = s3_result["size_bytes"]
            except Exception as s3_err:
                logger.warning(f"S3 upload failed (non-blocking): {s3_err}")

            # Métadonnées en PostgreSQL
            cursor = conn.execute(
                """INSERT INTO logs_uploaded (equipement, filename, s3_key, content_hash, size_bytes, nb_errors, nb_critiques, uploaded_by, parsed_errors)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (equipement, filename, s3_key, content_hash, size_bytes, nb_errors, nb_critiques, username, parsed_errors_str)
            )
            conn.execute(
                "INSERT INTO audit_log (username, action, details) VALUES (?, ?, ?)",
                (username, "Upload Log", f"Log '{filename}' S3:{s3_key or 'N/A'} ({nb_errors} erreurs)")
            )
            return {"ok": True, "id": cursor.lastrowid, "s3_key": s3_key,
                    "message": f"Log enregistré — {size_bytes} octets, {nb_errors} erreur(s), S3: {'ok' if s3_key else 'fallback'}"}
    except Exception as e:
        logger.error(f"Log upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/logs")
def list_logs(equipement: str = None, user: dict = Depends(_verify_token)):
    """Liste les logs uploadés (métadonnées depuis PostgreSQL)."""
    try:
        with get_db() as conn:
            if equipement:
                rows = conn.execute(
                    "SELECT id, equipement, filename, s3_key, size_bytes, nb_errors, nb_critiques, uploaded_by, uploaded_at FROM logs_uploaded WHERE equipement = ? ORDER BY uploaded_at DESC",
                    (equipement,)
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT id, equipement, filename, s3_key, size_bytes, nb_errors, nb_critiques, uploaded_by, uploaded_at FROM logs_uploaded ORDER BY uploaded_at DESC"
                ).fetchall()
            def _row(r):
                if isinstance(r, dict):
                    return {"id": r.get("id"), "equipement": r.get("equipement"), "filename": r.get("filename"), "s3_key": r.get("s3_key"), "size_bytes": r.get("size_bytes"), "nb_errors": r.get("nb_errors"), "nb_critiques": r.get("nb_critiques"), "uploaded_by": r.get("uploaded_by"), "uploaded_at": str(r.get("uploaded_at", ""))}
                return {"id": r[0], "equipement": r[1], "filename": r[2], "s3_key": r[3], "size_bytes": r[4], "nb_errors": r[5], "nb_critiques": r[6], "uploaded_by": r[7], "uploaded_at": str(r[8])}
            return [_row(r) for r in rows]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/logs/{log_id}")
def get_log(log_id: int, user: dict = Depends(_verify_token)):
    """Récupère le contenu d'un log depuis S3/MinIO."""
    try:
        with get_db() as conn:
            row = conn.execute("SELECT s3_key, equipement, filename, parsed_errors FROM logs_uploaded WHERE id = ?", (log_id,)).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Log non trouvé")
            # Support both dict (PG) and tuple (SQLite) rows
            if isinstance(row, dict):
                s3_key = row.get("s3_key", "")
                equipement = row.get("equipement", "")
                filename = row.get("filename", "")
            else:
                s3_key, equipement, filename = row[0], row[1], row[2]
            if not s3_key:
                import json as _json_nk
                _pe_nk = None
                if (isinstance(row, dict) and row.get("parsed_errors")) or (not isinstance(row, dict) and len(row) > 3 and row[3]):
                    _pe_nk_raw = row.get("parsed_errors") if isinstance(row, dict) else row[3]
                    try: _pe_nk = _json_nk.loads(_pe_nk_raw)
                    except: _pe_nk = None
                return {"id": log_id, "equipement": equipement, "filename": filename, "content": "", "parsed_errors": _pe_nk}
            try:
                from s3_storage import download_file
                content = download_file(s3_key)
                import json as _json_ret
                _pe = None
                if (isinstance(row, dict) and row.get("parsed_errors")) or (not isinstance(row, dict) and len(row) > 3 and row[3]):
                    _pe_raw = row.get("parsed_errors") if isinstance(row, dict) else row[3]
                    try: _pe = _json_ret.loads(_pe_raw)
                    except: _pe = None
                return {"id": log_id, "equipement": equipement, "filename": filename, "content": content or "", "s3_key": s3_key, "parsed_errors": _pe}
            except Exception:
                return {"id": log_id, "equipement": equipement, "filename": filename, "content": ""}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# AI INTEGRATION
# ==========================================

@app.post("/api/ai/analyze-diagnostic")
def analyze_diagnostic(body: dict, user: dict = Depends(_verify_token)):
    """Calls Gemini to diagnose a machine error code and log contexts."""
    try:
        from ai_engine import get_ai_suggestion, AI_AVAILABLE
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    if not AI_AVAILABLE:
        raise HTTPException(status_code=503, detail="L'IA n'est pas disponible. (Vérifiez GOOGLE_API_KEY).")

    machine = body.get("machine", "Équipement inconnu")
    code_erreur = body.get("code_erreur", "")
    message_erreur = body.get("message_erreur", "")
    log_context = body.get("log_context", "")

    try:
        result = get_ai_suggestion(code_erreur, message_erreur, machine, log_context=log_context)
        import json
        if isinstance(result, str):
            try:
                result = json.loads(result)
            except:
                return {"ok": True, "result": result}
        
        if result and isinstance(result, dict):
            # Map uppercase keys from ai_engine to lowercase keys expected by frontend
            return {"ok": True, "result": {
                "probleme": result.get("Probleme", result.get("probleme", "Non identifié")),
                "cause": result.get("Cause", result.get("cause", "À déterminer")),
                "solution": result.get("Solution", result.get("solution", "Analyse manuelle requise")),
                "prevention": result.get("Prevention", result.get("prevention", "Maintenance préventive recommandée")),
                "urgence": result.get("Urgence", result.get("urgence", "À évaluer")),
                "type": result.get("Type", result.get("type", "?")),
                "priorite": result.get("Priorite", result.get("priorite", "MOYENNE")),
                "confidence": result.get("Confidence_Score", result.get("confidence", 0)),
            }}
        return {"ok": True, "result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Diagnostic IA échoué: {e}")

@app.post("/api/ai/analyze-performance")
def analyze_performance(body: dict, user: dict = Depends(_verify_token)):
    """Calls Gemini to produce a detailed predictive maintenance report (v2)."""
    try:
        from ai_engine import _call_ia, clean_json_response, AI_AVAILABLE
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not AI_AVAILABLE:
        raise HTTPException(status_code=503, detail="L'IA n'est pas disponible.")

    kpis = body.get("kpis", {})
    sym = body.get("sym", "TND")

    # --- Fetch real per-machine data from DB ---
    machine_details = ""
    equip_detail = ""
    try:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT machine, COUNT(*) as nb, "
                "SUM(CASE WHEN type_intervention='Corrective' THEN 1 ELSE 0 END) as corr, "
                "SUM(CASE WHEN type_intervention ILIKE '%%r\u00e9ventive%%' THEN 1 ELSE 0 END) as prev, "
                "ROUND(AVG(duree_minutes)::numeric,1) as mttr_m, "
                "ROUND(SUM(cout)::numeric,0) as cout "
                "FROM interventions GROUP BY machine ORDER BY nb DESC LIMIT 20"
            ).fetchall()
            for r in rows:
                machine_details += f"  - {r['machine']}: {r['nb']} int ({r['corr']} corr, {r['prev']} prev), MTTR={r['mttr_m']}min, co\u00fbt={r['cout']} {sym}\n"
            eqs = conn.execute('SELECT "Nom","Client","Type","Statut","DateInstallation" FROM equipements ORDER BY "Nom" LIMIT 25').fetchall()
            for eq in eqs:
                equip_detail += f"  - {eq['Nom']} ({eq.get('Type','?')}) — {eq.get('Client','?')}, install\u00e9: {eq.get('DateInstallation','?')}, statut: {eq.get('Statut','?')}\n"
    except Exception as db_err:
        logger.warning(f"DB fetch for AI failed: {db_err}")

    risk_detail = ""
    for r in kpis.get("top_risques", []):
        risk_detail += f"  - {r.get('machine','?')}: risque={r.get('risque_panne_pct',0)}%, pi\u00e8ce={r.get('composant_a_risque','?')}, panne_dans={r.get('jours_avant_panne','?')}j, sant\u00e9={r.get('score_sante',0)}%\n"

    import datetime
    today = datetime.date.today()

    prompt = f"""Tu es Directeur du Service Technique d'une entreprise de maintenance d'\u00e9quipements d'imagerie m\u00e9dicale en Tunisie.
Analyse ces donn\u00e9es R\u00c9ELLES et produis un rapport pr\u00e9dictif d\u00e9taill\u00e9.

=== CHIFFRES DU PARC ===
- \u00c9quipements : {kpis.get('nb_equipements', 0)} | Interventions : {kpis.get('nb_interventions', 0)}
- Correctives : {kpis.get('interventions_correctives', 0)} | Pr\u00e9ventives : {kpis.get('interventions_preventives', 0)} | Calibrations : {kpis.get('interventions_calibration', 0)}
- Disponibilit\u00e9 : {kpis.get('disponibilite', 0)}% | MTBF : {kpis.get('mtbf', 0)}h | MTTR : {kpis.get('mttr', 0)}h
- Co\u00fbt total : {kpis.get('cout_total', 0)} {sym}

=== HISTORIQUE PAR MACHINE ===
{machine_details if machine_details else 'Non disponible'}

=== PR\u00c9DICTIONS IA ===
{risk_detail if risk_detail else 'Aucune'}

=== \u00c9QUIPEMENTS ===
{equip_detail if equip_detail else 'Non disponible'}

PRODUIS un rapport JSON STRICT :
{{{{
  "alertes_critiques": [
    {{{{
      "machine": "Nom (Client)",
      "score_sante": 41,
      "jours_avant_panne": 2,
      "nb_interventions": 19,
      "risque": "Risque concret",
      "action_immediate": "Action + pi\u00e8ces"
    }}}}
  ],
  "machines_stables": [
    {{{{
      "machine": "Nom (Client)",
      "score_sante": 84,
      "commentaire": "Pourquoi fiable"
    }}}}
  ],
  "plan_maintenance": [
    {{{{
      "jour": "Lundi {today.strftime('%d/%m')}",
      "cibles": "Machines",
      "action": "Action"
    }}}},
    {{{{
      "jour": "Mardi {(today + datetime.timedelta(days=1)).strftime('%d/%m')}",
      "cibles": "Machines",
      "action": "Action"
    }}}},
    {{{{
      "jour": "Mercredi {(today + datetime.timedelta(days=2)).strftime('%d/%m')}",
      "cibles": "Machines",
      "action": "Action"
    }}}}
  ],
  "estimation_couts": {{{{
    "cout_curatif_historique": {int(kpis.get('cout_total', 0))},
    "cout_preventif_propose": 0,
    "detail_preventif": "D\u00e9tail calcul",
    "gain_potentiel": 0,
    "ratio": "Pour 1 TND investi, X TND \u00e9conomis\u00e9s"
  }}}},
  "tendances": ["Tendance 1", "Tendance 2", "Tendance 3"],
  "conclusion": "Priorit\u00e9 absolue \u00e0..."
}}}}"""

    raw = _call_ia(prompt, timeout=90, is_json=True)
    if not raw:
        raise HTTPException(status_code=500, detail="L'IA n'a pas r\u00e9pondu.")
    result = clean_json_response(raw)
    return {"ok": True, "result": result}


@app.post("/api/ai/analyze-pieces")
def analyze_pieces(body: dict, user: dict = Depends(_verify_token)):
    """Calls Gemini to produce a spare parts purchase prediction report with dates and reasons."""
    try:
        from ai_engine import _call_ia, clean_json_response, AI_AVAILABLE
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not AI_AVAILABLE:
        raise HTTPException(status_code=503, detail="L'IA n'est pas disponible.")

    pieces_data = body.get("pieces", [])
    sym = body.get("sym", "TND")

    import datetime
    today = datetime.date.today()
    def fmt(d): return d.strftime("%d/%m/%Y")

    inventory_lines = ""
    for p in pieces_data:
        nom = p.get("designation","?"); ref = p.get("reference","?")
        stock = p.get("stock_actuel", 0); mini = p.get("stock_minimum", 1)
        prix = p.get("prix_unitaire", 0); four = p.get("fournisseur","N/A")
        equip = p.get("equipement_type","?")
        manquant = max(0, mini - stock + 1)
        if stock == 0: statut = "RUPTURE TOTALE"
        elif stock <= mini: statut = f"STOCK BAS (manque {manquant})"
        else: statut = f"OK (marge={stock-mini})"
        inventory_lines += f"  - {nom} ({ref}) | Equip: {equip} | Stock: {stock}/{mini} [{statut}] | {four} | {prix} {sym}\n"

    total_val = sum(p.get("stock_actuel",0)*p.get("prix_unitaire",0) for p in pieces_data)
    ruptures = sum(1 for p in pieces_data if p.get("stock_actuel",0)==0)
    bas = sum(1 for p in pieces_data if 0 < p.get("stock_actuel",0) <= p.get("stock_minimum",1))
    s1 = f"{fmt(today)} - {fmt(today+datetime.timedelta(days=6))}"
    s2 = f"{fmt(today+datetime.timedelta(days=7))} - {fmt(today+datetime.timedelta(days=13))}"
    s3 = f"{fmt(today+datetime.timedelta(days=14))} - {fmt(today+datetime.timedelta(days=20))}"
    d0=fmt(today); d3=fmt(today+datetime.timedelta(days=3)); d7=fmt(today+datetime.timedelta(days=7)); d14=fmt(today+datetime.timedelta(days=14))

    prompt = f"""Tu es Supply Chain Manager expert en pièces de rechange pour équipements de radiologie médicale (Tunisie).
Date du jour : {fmt(today)} | Inventaire : {len(pieces_data)} réf. | Valeur : {total_val:,.0f} {sym} | RUPTURES : {ruptures} | STOCK BAS : {bas}

INVENTAIRE :
{inventory_lines}

Génère un plan d'achat prévisionnel avec dates précises et raisons médicales/opérationnelles.
Réponds UNIQUEMENT en JSON valide (sans markdown, sans texte avant/après) :
{{
  "analyse_risque": "Synthèse 3-4 phrases sur pièces critiques, impact soins, capital immobilisé",
  "recommandations": [
    {{"piece": "Nom pièce", "reference": "REF", "raison": "Impact médical concret si non commandée (ex: arrêt scanner CT = patients sans diagnostic)", "action": "Commander immédiatement", "quantite": 2, "date_achat": "{d0}", "urgence": "critique", "cout_estime": 500}},
    {{"piece": "Nom pièce 2", "reference": "REF2", "raison": "Raison opérationnelle spécifique", "action": "Commander bientôt", "quantite": 1, "date_achat": "{d7}", "urgence": "haute", "cout_estime": 300}}
  ],
  "plan_achat": [
    {{"semaine": "S1 ({s1})", "pieces": ["pièce1"], "budget": 1200, "priorite": "Critique"}},
    {{"semaine": "S2 ({s2})", "pieces": ["pièce2"], "budget": 800, "priorite": "Haute"}},
    {{"semaine": "S3 ({s3})", "pieces": ["pièce3"], "budget": 500, "priorite": "Normale"}}
  ],
  "impact_budget": {{
    "cout_total_commande": 2500,
    "gain_potentiel": 8000,
    "ratio": "Pour 1 {sym} investi, 3.2 {sym} économisés en arrêts",
    "cout_indisponibilite_estime": 3000
  }},
  "tendances": ["Tendance 1 avec données concrètes", "Tendance 2", "Tendance 3"]
}}"""

    raw = _call_ia(prompt, timeout=90, is_json=True)
    if not raw:
        raise HTTPException(status_code=500, detail="L'IA n'a pas répondu.")
    result = clean_json_response(raw)
    return {"ok": True, "result": result}


@app.post("/api/ai/analyze-sav")
def analyze_sav(body: dict, user: dict = Depends(_verify_token)):
    """Comprehensive SAV/Interventions analysis using Gemini."""
    try:
        from ai_engine import _call_ia, clean_json_response, AI_AVAILABLE
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not AI_AVAILABLE:
        raise HTTPException(status_code=503, detail="L'IA n'est pas disponible.")

    sav_data = body.get("sav_data", {})
    sym = body.get("sym", "TND")

    prompt = f"""Tu es un expert en gestion de maintenance SAV pour équipements d'imagerie médicale en Tunisie.
Analyse ces données SAV RÉELLES et produis un rapport COMPLET et DÉTAILLÉ.

=== STATISTIQUES GLOBALES ===
- Total interventions : {sav_data.get('nb_total', 0)}
- Clôturées : {sav_data.get('nb_cloturees', 0)}
- En cours : {sav_data.get('nb_en_cours', 0)}
- Taux résolution : {sav_data.get('taux_resolution', 0)}%
- MTTR moyen : {sav_data.get('mttr_h', 0)}h
- Durée totale : {sav_data.get('duree_totale_h', 0)}h

=== RÉPARTITION PAR TYPE ===
- Correctives : {sav_data.get('nb_correctives', 0)}
- Préventives : {sav_data.get('nb_preventives', 0)}  
- Installations : {sav_data.get('nb_installations', 0)}
- Ratio correctif : {sav_data.get('ratio_correctif_pct', 0)}%

=== COÛTS ===
- Coût total interventions : {sav_data.get('cout_interventions', 0)} {sym}
- Coût pièces : {sav_data.get('cout_pieces', 0)} {sym}
- Coût total : {sav_data.get('cout_total', 0)} {sym}
- Coût moyen/intervention : {sav_data.get('cout_moyen', 0)} {sym}

=== PERFORMANCE ÉQUIPE (par technicien) ===
{sav_data.get('tech_details', 'Non disponible')}

=== DÉTAIL DES INTERVENTIONS RÉCENTES ===
{sav_data.get('interventions_detail', 'Non disponible')}

=== MACHINES LES PLUS INTERVENUES ===
{sav_data.get('machines_detail', 'Non disponible')}

=== CLIENTS ===
{sav_data.get('clients_detail', 'Non disponible')}

IMPORTANT: Analyse en profondeur et produis un JSON STRICT avec cette structure exacte :
{{{{
  "analyse": "Résumé exécutif complet de la situation SAV (3-5 phrases détaillées)",
  "score_global": 75,
  "points_forts": [
    "Point fort 1 détaillé avec chiffres",
    "Point fort 2 détaillé avec chiffres",
    "Point fort 3 détaillé avec chiffres"
  ],
  "points_faibles": [
    "Point faible 1 détaillé avec chiffres",
    "Point faible 2 détaillé avec chiffres", 
    "Point faible 3 détaillé avec chiffres"
  ],
  "recommandations": [
    {{{{
      "titre": "Titre recommandation",
      "description": "Description détaillée de l'action à entreprendre",
      "impact": "HAUT"
    }}}},
    {{{{
      "titre": "Titre recommandation 2",
      "description": "Description détaillée",
      "impact": "MOYEN"
    }}}},
    {{{{
      "titre": "Titre recommandation 3",
      "description": "Description détaillée",
      "impact": "BAS"
    }}}}
  ],
  "performance_equipe": [
    {{{{
      "technicien": "Nom",
      "evaluation": "Excellent/Bon/À améliorer",
      "commentaire": "Commentaire détaillé sur ses performances"
    }}}}
  ],
  "analyse_couts": {{{{
    "verdict": "Maîtrisés/Élevés/Critiques",
    "detail": "Analyse détaillée des coûts",
    "economie_possible": "Estimation d'économie possible et comment"
  }}}},
  "tendances": [
    "Tendance 1 observée",
    "Tendance 2 observée",
    "Tendance 3 observée"
  ],
  "priorites_immediates": [
    "Action prioritaire 1",
    "Action prioritaire 2"
  ]
}}}}"""

    raw = _call_ia(prompt, timeout=90, is_json=True)
    if not raw:
        raise HTTPException(status_code=500, detail="L'IA n'a pas répondu.")
    result = clean_json_response(raw)
    return {"ok": True, "result": result}


# ==========================================
# ADMIN — Utilisateurs
# ==========================================

@app.get("/api/admin/users")
def get_users(user: dict = Depends(_verify_token)):
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, username, nom_complet, role, client, actif, created_at, last_login FROM utilisateurs ORDER BY id"
        ).fetchall()
    return [dict(r) for r in rows]


@app.post("/api/admin/users")
def create_user(body: dict, user: dict = Depends(_verify_token)):
    hashed = bcrypt.hashpw(body["password"].encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    with get_db() as conn:
        conn.execute(
            "INSERT INTO utilisateurs (username, password_hash, nom_complet, role, client, actif) VALUES (?, ?, ?, ?, ?, 1)",
            (body["username"], hashed, body.get("nom_complet", ""), body.get("role", "Lecteur"), body.get("client", ""))
        )
    return {"ok": True}


@app.put("/api/admin/users/{user_id}")
def update_user(user_id: int, body: dict, user: dict = Depends(_verify_token)):
    fields = []
    params = []
    for f in ["nom_complet", "role", "client", "actif"]:
        if f in body:
            fields.append(f"{f} = ?")
            params.append(body[f])
    if "password" in body and body["password"]:
        fields.append("password_hash = ?")
        params.append(bcrypt.hashpw(body["password"].encode("utf-8"), bcrypt.gensalt()).decode("utf-8"))
    if fields:
        params.append(user_id)
        with get_db() as conn:
            conn.execute(f"UPDATE utilisateurs SET {', '.join(fields)} WHERE id = ?", params)
    return {"ok": True}


@app.delete("/api/admin/users/{user_id}")
def delete_user(user_id: int, user: dict = Depends(_verify_token)):
    with get_db() as conn:
        conn.execute("DELETE FROM utilisateurs WHERE id = ?", (user_id,))
    return {"ok": True}


# ==========================================
# AUDIT LOG
# ==========================================

@app.get("/api/audit")
def get_audit_log(limit: int = 100, user: dict = Depends(_verify_token)):
    return _df_to_records(lire_audit(limit=limit))


# ==========================================
# NOTIFICATIONS
# ==========================================

@app.get("/api/notifications")
def get_notifications(
    destination: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    """Liste les notifications. Destination auto-detectée selon le rôle."""
    role = user.get("role", "")
    nom = (user.get("nom") or "").strip()
    if destination is None:
        destination = "technicien" if role == "Technicien" else "gestionnaire"
    df = lire_notifications_pieces(destination=destination)
    # Pour les techniciens : filtrer par leur nom
    if role == "Technicien" and nom and not df.empty:
        from db_engine import read_sql
        df = df[df["technicien"].fillna("").str.lower().apply(
            lambda t: all(w in t for w in nom.lower().split() if len(w) > 1)
        )]
    return _df_to_records(df)


@app.get("/api/notifications/count")
def get_notification_count(
    destination: Optional[str] = None,
    user: dict = Depends(_verify_token),
):
    """Compte les notifications non lues. Destination auto-detectée selon le rôle."""
    role = user.get("role", "")
    nom = (user.get("nom") or "").strip()
    if destination is None:
        destination = "technicien" if role == "Technicien" else "gestionnaire"
    count = compter_notifications_non_lues(destination, technicien=nom if role == "Technicien" else None)
    return {"count": count}


@app.patch("/api/notifications/{notif_id}/read")
def mark_notification_read(notif_id: int, user: dict = Depends(_verify_token)):
    """Marque une notification comme lue."""
    marquer_notification_lue(notif_id)
    return {"ok": True}


@app.patch("/api/notifications/{notif_id}/done")
def mark_notification_done(notif_id: int, user: dict = Depends(_verify_token)):
    """Marque une notification comme traitée."""
    marquer_notification_traitee(notif_id)
    return {"ok": True}


# ==========================================
# SETTINGS / CONFIG
# ==========================================

@app.get("/api/settings")
def get_settings(user: dict = Depends(_verify_token)):
    keys = [
        "nom_organisation", "logo_path", "langue", "theme",
        "taux_horaire_technicien", "telegram_token", "telegram_chat_id",
        "gemini_api_key", "role_permissions",
    ]
    try:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT cle, valeur FROM config_client WHERE cle = ANY(%s)",
                (keys,)
            ).fetchall()
            result = {k: "" for k in keys}
            for row in rows:
                result[row["cle"]] = row["valeur"] or ""
            return result
    except Exception:
        # Fallback: chercher clé par clé
        result = {}
        for k in keys:
            result[k] = get_config(k, "")
        return result


@app.put("/api/settings")
def update_settings(body: dict, user: dict = Depends(_verify_token)):
    try:
        with get_db() as conn:
            for k, v in body.items():
                conn.execute(
                    """
                    INSERT INTO config_client (cle, valeur) VALUES (%s, %s)
                    ON CONFLICT (cle) DO UPDATE SET valeur = EXCLUDED.valeur
                    """,
                    (k, str(v))
                )
        return {"ok": True}
    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=f"Erreur sauvegarde config: {e}\n{traceback.format_exc()}")


# ==========================================
# CLIENTS (derived from equipements)
# ==========================================

@app.get("/api/clients")
def get_clients(user: dict = Depends(_verify_token)):
    """Derive client list from equipements + contrats + interventions."""
    df_eq = lire_equipements()
    df_int = lire_interventions()
    df_contrats = lire_contrats()

    clients = {}
    if not df_eq.empty and "Client" in df_eq.columns:
        for client_name in df_eq["Client"].unique():
            if not client_name:
                continue
            eq_client = df_eq[df_eq["Client"] == client_name]
            nb_eq = len(eq_client)
            # Health score
            nb_hs = len(eq_client[eq_client["Statut"].isin(["Hors Service", "Critique"])]) if "Statut" in eq_client.columns else 0
            score = max(0, round(((nb_eq - nb_hs) / nb_eq) * 100)) if nb_eq > 0 else 100
            # Interventions count
            nb_int = 0
            if not df_int.empty and "machine" in df_int.columns:
                machines = eq_client["Nom"].tolist() if "Nom" in eq_client.columns else []
                nb_int = len(df_int[df_int["machine"].isin(machines)])
            # Active contracts
            nb_contrats = 0
            if not df_contrats.empty and "client" in df_contrats.columns:
                nb_contrats = len(df_contrats[df_contrats["client"] == client_name])

            clients[client_name] = {
                "nom": client_name,
                "nb_equipements": nb_eq,
                "nb_interventions": nb_int,
                "nb_contrats": nb_contrats,
                "score_sante": score,
            }

    return list(clients.values())


# ==========================================
# LOGS / S3 MANAGEMENT
# ==========================================

try:
    import s3_storage
except ImportError:
    s3_storage = None


@app.get("/api/logs")
def api_list_logs(machine: Optional[str] = None, user=Depends(_verify_token)):
    """List all logs from S3, optionally filtered by machine name."""
    if not s3_storage or not s3_storage.S3_AVAILABLE:
        s3_storage and s3_storage._init_s3()
    if not s3_storage or not s3_storage.S3_AVAILABLE:
        return []

    prefix = "logs/"
    if machine:
        # Search across all date folders for this machine
        all_files = s3_storage.list_files(prefix)
        return [f for f in all_files if f"/{machine.replace(' ', '_')}/" in f["key"] or machine.replace(' ', '_') in f["key"]]
    return s3_storage.list_files(prefix)


@app.delete("/api/logs")
def api_delete_log(key: str = Query(..., description="S3 key of the log to delete"), user=Depends(_verify_token)):
    """Delete a specific log file from S3 by its key."""
    if not s3_storage:
        raise HTTPException(status_code=503, detail="S3 storage not available")
    s3_storage._init_s3()
    if not s3_storage.S3_AVAILABLE:
        raise HTTPException(status_code=503, detail="S3 storage not connected")

    success = s3_storage.delete_file(key)
    if success:
        return {"ok": True, "message": f"Log supprimé: {key}"}
    raise HTTPException(status_code=500, detail="Échec de la suppression du log")


@app.delete("/api/logs/machine/{machine_name}")
def api_delete_machine_logs(machine_name: str, user=Depends(_verify_token)):
    """Delete ALL log files for a given machine from S3."""
    if not s3_storage:
        raise HTTPException(status_code=503, detail="S3 storage not available")
    s3_storage._init_s3()
    if not s3_storage.S3_AVAILABLE:
        raise HTTPException(status_code=503, detail="S3 storage not connected")

    # Find all logs matching this machine
    all_files = s3_storage.list_files("logs/")
    machine_key = machine_name.replace(' ', '_')
    to_delete = [f for f in all_files if f"/{machine_key}/" in f["key"] or machine_key in f["key"]]

    deleted = 0
    for f in to_delete:
        if s3_storage.delete_file(f["key"]):
            deleted += 1

    return {"ok": True, "deleted": deleted, "message": f"{deleted} log(s) supprimé(s) pour {machine_name}"}




# ==========================================
# PDF REPORT GENERATION (server-side fpdf2)
# ==========================================

def _sanitize(text):
    if not text:
        return ""
    text = str(text)
    # Replace specific chars with ASCII equivalents
    text = text.replace(chr(0x2014), " - ")  # em dash
    text = text.replace(chr(0x2013), " - ")  # en dash
    text = text.replace(chr(0x202F), " ")     # narrow no-break space
    text = text.replace(chr(0x00A0), " ")     # no-break space
    text = text.replace(chr(0x2022), "-")     # bullet
    text = text.replace(chr(0x2019), "'")    # right single quote
    text = text.replace(chr(0x2018), "'")    # left single quote
    text = text.replace(chr(0x201C), '"')    # left double quote
    text = text.replace(chr(0x201D), '"')    # right double quote
    text = text.replace(chr(0x2026), "...")   # ellipsis
    text = text.replace(chr(0x20AC), "EUR")   # euro sign
    # Final fallback: encode to Latin-1, unknown chars become ?
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _fmt_number(n):
    try:
        val = int(round(float(n)))
        s = str(abs(val))
        result = ""
        for i, c in enumerate(reversed(s)):
            if i > 0 and i % 3 == 0:
                result = " " + result
            result = c + result
        return ("-" if val < 0 else "") + result
    except Exception:
        return str(n)


class SaviaPDF(FPDF):
    """FPDF subclass with auto-repeated compact header on every page."""
    _savia_logo   = None    # path
    _client_logo  = None    # BytesIO (seekable)
    _company_name = ''
    _company_sub  = ''
    _report_title = ''     # centered between logos
    HEADER_H = 26           # header height mm

    def header(self):
        from io import BytesIO
        H = self.HEADER_H
        y0 = 5
        W  = self.w - 16   # 8mm each side

        # ── SAVIA logo (left) ──────────────────────────────────────
        savia_w = 0
        if self._savia_logo and os.path.exists(self._savia_logo):
            try:
                self.image(self._savia_logo, x=8, y=y0, h=H - 4)
                savia_w = 24
            except Exception:
                savia_w = 0

        # ── Client logo (right) ───────────────────────────────────
        client_w = 0
        if self._client_logo:
            try:
                self._client_logo.seek(0)
                self.image(self._client_logo, x=self.w - 8 - 28, y=y0, h=H - 4)
                client_w = 30
            except Exception:
                client_w = 0

        # ── Center zone: report title + company name ───────────────
        cx = 8 + savia_w + 2
        cw = W - savia_w - client_w - 4

        # Report title (top, bold, centered between logos)
        if self._report_title:
            self.set_xy(cx, y0 + 1)
            self.set_font('Helvetica', 'B', 11)
            self.set_text_color(30, 40, 55)
            self.cell(cw, 7, _sanitize(self._report_title[:70]), align='C')

        # Company name (below title, smaller)
        if self._company_name:
            y_cn = y0 + 9 if self._report_title else y0 + 4
            self.set_xy(cx, y_cn)
            self.set_font('Helvetica', 'B', 9)
            self.set_text_color(1, 180, 188)
            self.cell(cw, 6, _sanitize(self._company_name[:55]), align='C')
            if self._company_sub:
                self.set_xy(cx, y_cn + 6)
                self.set_font('Helvetica', '', 7)
                self.set_text_color(130, 145, 160)
                self.cell(cw, 4.5, _sanitize(self._company_sub[:90]), align='C')

        # ── Separator line ─────────────────────────────────────────
        sep = y0 + H
        self.set_fill_color(1, 180, 188)
        self.rect(8, sep, self.w - 16, 0.8, style='F')
        self.set_y(sep + 3)
        self.set_text_color(40, 50, 65)

    def set_header_data(self, savia_logo, client_logo_bytes, company_name, company_sub, report_title=""):
        self._savia_logo   = savia_logo
        self._client_logo  = client_logo_bytes
        self._company_name = company_name
        self._company_sub  = company_sub
        self._report_title = report_title



class PdfRequest(BaseModel):
    title: str = "Rapport SAVIA"
    subtitle: str = ""
    filename: str = "rapport"
    company_name: str = "SAVIA"
    company_logo: str = ""
    kpis: list = []
    head: list = []
    rows: list = []
    type_data: list = []
    table_title: str = ""
    is_ai_report: bool = False
    ai_content: str = ""


@app.post("/api/reports/generate-pdf")
def generate_pdf_report(data: PdfRequest, user: dict = Depends(_verify_token)):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from io import BytesIO
    from fastapi.responses import Response
    import json

    SAVIA_LOGO = "/app/logo-savia.png"
    try:
        from io import BytesIO as _BytesIO
        import base64 as _b64
        import urllib.request as _ur

        orientation = "P" if data.is_ai_report else "L"
        pdf = SaviaPDF(orientation=orientation, unit="mm", format="A4")

        # ── Resolve client logo (URL or base64) ──────────────
        _client_logo_io = None
        if data.company_logo:
            try:
                clogo = data.company_logo.strip()
                if clogo.startswith("data:"):
                    _b64_part = clogo.split(",", 1)[1] if "," in clogo else clogo
                    _client_logo_io = _BytesIO(_b64.b64decode(_b64_part))
                elif clogo.startswith("http"):
                    req_ = _ur.Request(clogo, headers={"User-Agent": "Mozilla/5.0"})
                    with _ur.urlopen(req_, timeout=6) as _r:
                        _client_logo_io = _BytesIO(_r.read())
            except Exception as _e:
                logger.warning(f"Client logo error: {_e}")

        pdf.set_header_data(
            SAVIA_LOGO, _client_logo_io,
            data.company_name if data.company_name != "SAVIA" else "",
            "Systeme Intelligent de Controle et de Gestion",  # Always fixed - footer has date/name
            report_title=data.title if data.title and data.title != "Rapport SAVIA" else ""
        )
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_top_margin(pdf.HEADER_H + 10)  # content starts below header
        pdf.add_page()
        page_w = pdf.w

        # Title is now shown in header (between logos)


        # KPIs
        if data.kpis:
            box_w, box_h, margin_ = 64, 16, 5
            kpi_y = pdf.get_y()
            for i, kpi in enumerate(data.kpis[:4]):
                kx = 10 + i * (box_w + margin_)
                color = kpi.get("color", [15, 118, 110])
                r1,g1,b1 = int(color[0]),int(color[1]),int(color[2])
                lc = [min(r1+215,255), min(g1+215,255), min(b1+215,255)]
                pdf.set_fill_color(*lc)
                pdf.set_draw_color(r1,g1,b1)
                pdf.set_line_width(0.4)
                pdf.rect(kx, kpi_y, box_w, box_h, style="FD")
                # Top accent bar (Dopely palette solid)
                pdf.set_fill_color(r1,g1,b1)
                pdf.rect(kx, kpi_y, box_w, 3, style="F")
                # Small white round dot on top bar
                pdf.set_fill_color(255, 255, 255)
                pdf.ellipse(kx + box_w/2 - 1.5, kpi_y + 0.3, 3, 2.4, style="F")
                # Value
                pdf.set_xy(kx, kpi_y + 3)
                pdf.set_font("Helvetica", "B", 13)
                vr = max(30, min(r1-30, 180)); vg = max(30, min(g1-20, 120)); vb = max(30, min(b1-20, 150))
                pdf.set_text_color(vr, vg, vb)
                pdf.cell(box_w, 8, _sanitize(str(kpi.get("val", ""))), align="C")
                # Label
                pdf.set_xy(kx, kpi_y + 11)
                pdf.set_font("Helvetica", "", 6.5)
                pdf.set_text_color(90, 100, 115)
                pdf.cell(box_w, 4, _sanitize(str(kpi.get("label", ""))), align="C")
            pdf.set_y(kpi_y + box_h + 6)
            pdf.set_text_color(0, 0, 0)

        # Type distribution table
        if data.type_data:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(50, 70, 90)
            pdf.cell(0, 6, "Repartition par type", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_fill_color(1, 180, 188)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(100, 7, "Type", border=1, fill=True)
            pdf.cell(30, 7, "Nombre", border=1, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 8)
            for idx, row in enumerate(data.type_data):
                fill = idx % 2 == 0
                if fill: pdf.set_fill_color(225, 250, 251)
                else: pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(30, 40, 60)
                pdf.cell(100, 6, _sanitize(str(row[0])) if row else "", border=1, fill=fill)
                pdf.cell(30, 6, str(row[1]) if len(row) > 1 else "", border=1, fill=fill, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(4)

        # AI Report mode - DOPELY PALETTE + ROUND BULLETS
        if data.is_ai_report and data.ai_content:
            try: ai = json.loads(data.ai_content)
            except Exception: ai = {"summary": data.ai_content}

            W = page_w - 20

            def sec_hdr(lbl, bg, fg=None):
                """Section header: solid color band + small white circle icon"""
                if fg is None: fg = [255,255,255]
                if pdf.get_y() > pdf.h - 40: pdf.add_page()
                yh = pdf.get_y()
                pdf.set_fill_color(bg[0], bg[1], bg[2])
                pdf.rect(10, yh, W, 8.5, style="F")
                # Label (no white dot)
                pdf.set_xy(14, yh + 1.8)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_text_color(fg[0], fg[1], fg[2])
                pdf.cell(W - 9, 5.5, _sanitize(lbl.upper()))
                pdf.set_y(yh + 11)
                pdf.set_text_color(40, 50, 65)

            def body_item(txt, bullet_color=None):
                """List item with round colored bullet"""
                if not txt: return
                if pdf.get_y() > pdf.h - 20: pdf.add_page()
                yi = pdf.get_y()
                if bullet_color:
                    # ROUND bullet (ellipse)
                    pdf.set_fill_color(bullet_color[0], bullet_color[1], bullet_color[2])
                    pdf.ellipse(13.5, yi + 2.2, 2.5, 2.5, style="F")
                    pdf.set_xy(18, yi)
                else:
                    pdf.set_x(13)
                pdf.set_font("Helvetica", "", 8.5)
                pdf.set_text_color(40, 50, 65)
                ww = W - 9 if bullet_color else W - 3
                pdf.multi_cell(ww, 4.8, _sanitize(str(txt)[:300]))
                pdf.ln(0.5)

            def add_sec(lbl, items, bg, fg=None):
                if not items: return
                sec_hdr(lbl, bg, fg)
                pdf.set_font("Helvetica", "", 8.5)
                for it in items:
                    txt = it if isinstance(it, str) else it.get("action", it.get("machine", str(it))) if isinstance(it, dict) else str(it)
                    body_item(txt, bg)
                pdf.ln(3)

            # ── Score global ─────────────────────────────────────
            score = ai.get("score_global")
            if score is not None:
                sc = int(score)
                if sc >= 70:
                    s_bg = [95,165,90]     # GREEN  #5FA55A
                elif sc >= 40:
                    s_bg = [250,137,37]    # ORANGE #FA8925
                else:
                    s_bg = [250,84,87]     # CORAL  #FA5457
                y_sc = pdf.get_y()
                # White bg + colored border + colored left bar
                pdf.set_fill_color(255, 255, 255)
                pdf.set_draw_color(s_bg[0], s_bg[1], s_bg[2])
                pdf.set_line_width(0.8)
                pdf.rect(10, y_sc, W, 16, style="FD")
                pdf.set_fill_color(s_bg[0], s_bg[1], s_bg[2])
                pdf.rect(10, y_sc, 5, 16, style="F")
                pdf.set_xy(18, y_sc + 2)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_text_color(40, 50, 65)
                pdf.cell(60, 5.5, "Score : " + str(sc) + "/100")
                slabel = "Excellent" if sc>=70 else "Satisfaisant" if sc>=40 else "A ameliorer"
                pdf.set_xy(80, y_sc + 2)
                pdf.set_font("Helvetica", "B", 8)
                pdf.set_text_color(s_bg[0], s_bg[1], s_bg[2])
                pdf.cell(W - 72, 5.5, slabel)
                pdf.set_xy(18, y_sc + 9)
                pdf.set_font("Helvetica", "", 7.5)
                pdf.set_text_color(80, 95, 110)
                pdf.cell(W - 8, 5, "Score global de performance de la periode")
                pdf.set_y(y_sc + 19)

            # ── Résumé Exécutif ── TEAL #01B4BC
            analyse = ai.get("analyse") or ai.get("summary")
            if analyse:
                sec_hdr("Resume Executif", [1,180,188])
                pdf.set_font("Helvetica", "", 8.5)
                pdf.set_text_color(40, 50, 65)
                pdf.set_x(13)
                pdf.multi_cell(W - 3, 5, _sanitize(str(analyse)[:2500]))
                pdf.ln(4)

            # ── Points Forts ── GREEN #5FA55A
            add_sec("Points Forts",     ai.get("points_forts", []),     [95,165,90])
            # ── Points Faibles ── CORAL #FA5457
            add_sec("Points Faibles",   ai.get("points_faibles", []),   [250,84,87])

            # ── Recommandations ── TEAL #01B4BC
            recs = ai.get("recommandations", [])
            recs_c = [r if isinstance(r, str) else r.get("action", str(r)) for r in recs]
            add_sec("Recommandations",  recs_c,                         [1,180,188])
            # ── Alertes Critiques ── CORAL #FA5457
            add_sec("Alertes Critiques",ai.get("alertes_critiques", []),[250,84,87])
            # ── Tendances ── TEAL #01B4BC
            add_sec("Tendances",        ai.get("tendances", []),        [1,180,188])

            # ── Performance Equipe ── YELLOW #F6D51F (dark text)
            perf = ai.get("performance_equipe", [])
            if perf:
                sec_hdr("Performance Equipe", [246,213,31], fg=[50,40,0])
                pdf.set_font("Helvetica", "", 8.5)
                pdf.set_text_color(40, 50, 65)
                for pe in perf:
                    if isinstance(pe, dict):
                        nm_ = pe.get("technicien", pe.get("nom", ""))
                        sc_ = pe.get("score", pe.get("note", ""))
                        dt_ = pe.get("detail", pe.get("commentaire", ""))
                        txt = _sanitize(str(nm_))+" : "+str(sc_)+(" - "+str(dt_) if dt_ else "")
                    else: txt = str(pe)
                    body_item(txt, [246,213,31])
                pdf.ln(3)

            # ── Analyse Financière ── ORANGE #FA8925
            couts = ai.get("analyse_couts")
            if couts and isinstance(couts, dict):
                sec_hdr("Analyse Financiere", [250,137,37])
                pdf.set_font("Helvetica", "", 8.5)
                pdf.set_text_color(40, 50, 65)
                for k_, v_ in couts.items():
                    if v_: body_item(str(k_)+" : "+str(v_), [250,137,37])
                pdf.ln(3)

            # ── Priorités Immédiates ── ORANGE #FA8925
            add_sec("Priorites Immediates", ai.get("priorites_immediates",[]),[250,137,37])

            # ── Conclusion ── TEAL #01B4BC
            conclusion = ai.get("conclusion")
            if conclusion:
                sec_hdr("Conclusion", [1,180,188])
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(50, 62, 78)
                pdf.set_x(13)
                pdf.multi_cell(W - 3, 5, _sanitize(str(conclusion)[:2500]))
        # Standard table
        elif data.head and data.rows:
            if pdf.get_y() > pdf.h - 45: pdf.add_page()
            if data.table_title:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(50, 70, 90)
                pdf.cell(0, 7, _sanitize(data.table_title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            n_cols = len(data.head)
            total_w = page_w - 20
            last_w = 28
            def_w = (total_w - last_w) / max(n_cols - 1, 1)
            col_w = [def_w] * (n_cols - 1) + [last_w]
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_fill_color(1, 180, 188)
            pdf.set_text_color(255, 255, 255)
            for i, h in enumerate(data.head):
                pdf.cell(col_w[i], 8, _sanitize(str(h)[:20]), border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 7.5)
            for row_idx, row in enumerate(data.rows):
                if pdf.get_y() > pdf.h - 15:
                    pdf.add_page()
                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_fill_color(15, 118, 110)
                    pdf.set_text_color(255, 255, 255)
                    for i, h in enumerate(data.head):
                        pdf.cell(col_w[i], 8, _sanitize(str(h)[:20]), border=1, fill=True, align="C")
                    pdf.ln()
                    pdf.set_font("Helvetica", "", 7.5)
                fill = row_idx % 2 == 0
                if fill: pdf.set_fill_color(244, 252, 251)
                else: pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(25, 35, 55)
                for i, cell in enumerate(row[:n_cols]):
                    val = _sanitize(_fmt_number(cell)) if i == n_cols - 1 else _sanitize(str(cell)[:25]) if cell else "-"
                    align = "R" if i == n_cols - 1 else "L"
                    pdf.cell(col_w[i], 6.5, val, border=1, fill=fill, align=align)
                pdf.ln()

        # ── CRITICAL: disable auto-page-break before footer loop ──────────────
        # cell() at y=h-10=287mm exceeds auto-break threshold (h-15=282mm)
        # → triggers unwanted new page with "Genere par" at top
        pdf.set_auto_page_break(auto=False)

        # Also remove last page if only header drawn (extra safety)
        try:
            _EMPTY_THRESHOLD = 46
            _last_y = pdf.get_y()
            _n_pages = len(pdf.pages)
            logger.info(f"PDF: {_n_pages} pages, last_y={_last_y:.1f}mm")
            if _last_y <= _EMPTY_THRESHOLD and _n_pages > 1:
                _last_pg = max(pdf.pages.keys())
                del pdf.pages[_last_pg]
                pdf.page = _last_pg - 1
                logger.info(f"Removed empty last page #{_last_pg}")
        except Exception as _ep:
            logger.warning(f"Empty page removal: {_ep}")

        # Footer on all pages (auto-break already disabled above)
        total_pages = len(pdf.pages)
        now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        for pg in range(1, total_pages + 1):
            pdf.page = pg
            pdf.set_xy(10, pdf.h - 11)
            pdf.set_draw_color(200, 205, 220)
            pdf.set_line_width(0.3)
            pdf.line(10, pdf.h - 11, page_w - 10, pdf.h - 11)
            pdf.set_xy(10, pdf.h - 9)
            pdf.set_font("Helvetica", "", 7)
            pdf.set_text_color(160, 170, 190)
            pdf.cell(page_w - 40, 5, _sanitize(f"Genere par {data.company_name} - {now_str}"), align="L")
            pdf.cell(30, 5, f"Page {pg} / {total_pages}", align="R")

        pdf_bytes = bytes(pdf.output())
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{data.filename}.pdf"',
                "Content-Length": str(len(pdf_bytes)),
            }
        )
    except Exception as e:
        logging.error(f"PDF generation error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

# ==========================================
# ENTRY POINT
# ==========================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
