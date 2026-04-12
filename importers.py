# ==========================================
# 📥 IMPORTEURS (EXCEL / CSV / PDF)
# ==========================================
import pandas as pd
import fitz  # PyMuPDF
import time
import streamlit as st
from database import ajouter_codes_batch
from ai_engine import AI_AVAILABLE, extraire_erreurs_texte, extraire_erreurs_image


def importer_donnees_structurees(fichier_upload, excel_path):
    """
    Importe des données depuis un fichier CSV ou Excel structuré.
    Gère les fichiers multi-feuilles (knowledge_base.xlsx) et les fichiers simples.
    Retourne le nombre de codes importés.
    """
    try:
        if fichier_upload.name.endswith(".csv"):
            df = pd.read_csv(fichier_upload)
            return _import_single_sheet(df, excel_path)

        # Excel : vérifier si c'est un fichier multi-feuilles SIC
        xls = pd.ExcelFile(fichier_upload)
        sheets = [s.upper() for s in xls.sheet_names]

        # Détecter le format SIC (knowledge_base.xlsx)
        has_codes = any("CODE" in s for s in sheets)
        has_solutions = any("SOLUTION" in s for s in sheets)

        if has_codes or has_solutions:
            return _import_sic_format(xls, excel_path)
        else:
            # Format générique : lire la première feuille
            df = pd.read_excel(fichier_upload)
            return _import_single_sheet(df, excel_path)

    except Exception as e:
        st.error(f"Erreur import: {e}")
        return 0


def _import_sic_format(xls, excel_path):
    """Importe un fichier au format SIC (multi-feuilles CODES_HEXA + SOLUTIONS_TEXTE)."""
    total = 0
    rows_hex = []
    rows_txt = []

    # Chercher la feuille CODES
    for sheet_name in xls.sheet_names:
        if "CODE" in sheet_name.upper():
            df_codes = pd.read_excel(xls, sheet_name)
            for _, row in df_codes.iterrows():
                code = str(row.get("Code", "")).strip()
                if len(code) > 1:
                    rows_hex.append({
                        "Code": code,
                        "Message": str(row.get("Message", "")),
                        "Niveau": str(row.get("Niveau", "ATTENTION")),
                        "Type": str(row.get("Type", "Import")),
                    })

    # Chercher la feuille SOLUTIONS
    for sheet_name in xls.sheet_names:
        if "SOLUTION" in sheet_name.upper():
            df_sol = pd.read_excel(xls, sheet_name)
            for _, row in df_sol.iterrows():
                mot_cle = str(row.get("Mot_Cle", "")).strip()
                if len(mot_cle) > 1:
                    rows_txt.append({
                        "Mot_Cle": mot_cle,
                        "Type": str(row.get("Type", "Import")),
                        "Priorite": str(row.get("Priorite", "MOYENNE")),
                        "Cause": str(row.get("Cause", "")),
                        "Solution": str(row.get("Solution", "")),
                    })

    if rows_hex or rows_txt:
        ajouter_codes_batch(excel_path, rows_hex, rows_txt)
        total = max(len(rows_hex), len(rows_txt))

    return total


def _import_single_sheet(df, excel_path):
    """Importe depuis un DataFrame à une seule feuille (CSV ou Excel simple)."""
    try:
        # Mapping intelligent des colonnes
        map_keys = {
            "Code": ["code", "id", "hex", "error_code", "err", "mot_cle"],
            "Message": ["message", "desc", "description", "msg", "erreur"],
            "Type_Reel": ["type", "cat", "catégorie", "domaine", "category"],
            "Cause": ["cause", "origin", "origine", "raison"],
            "Solution": ["sol", "action", "fix", "resolution", "solution"],
            "Priorite": ["prio", "level", "priorité", "priority", "niveau"],
        }

        df_final = pd.DataFrame()
        defaults = {
            "Code": "?", "Message": "Importé", "Type_Reel": "Import",
            "Cause": "?", "Solution": "Voir documentation", "Priorite": "MOYENNE",
        }

        for target, keywords in map_keys.items():
            matching_cols = [
                c for c in df.columns if any(k in c.lower() for k in keywords)
            ]
            if matching_cols:
                if len(matching_cols) > 1:
                    df_final[target] = df[matching_cols].astype(str).agg(" - ".join, axis=1)
                else:
                    df_final[target] = df[matching_cols[0]].astype(str)
            else:
                df_final[target] = defaults.get(target, "")

        # Filtrer les lignes vides
        df_final = df_final[df_final["Code"].str.len() > 1]

        rows_hex = []
        rows_txt = []
        for _, row in df_final.iterrows():
            code = str(row["Code"]).strip()
            rows_hex.append({
                "Code": code,
                "Message": row["Message"],
                "Niveau": "ATTENTION",
                "Type": str(row["Type_Reel"]),
            })
            rows_txt.append({
                "Mot_Cle": code,
                "Type": str(row["Type_Reel"]),
                "Priorite": row["Priorite"],
                "Cause": row["Cause"],
                "Solution": row["Solution"],
            })

        if rows_hex:
            ajouter_codes_batch(excel_path, rows_hex, rows_txt)

        return len(rows_hex)

    except Exception as e:
        st.error(f"Erreur import: {e}")
        return 0


def importer_pdf(uploaded_pdf, page_debut, page_fin, excel_path):
    """
    Importe des codes d'erreur depuis un PDF technique.
    Utilise l'IA pour extraire les informations des pages.
    Retourne le nombre de codes extraits.
    """
    if not AI_AVAILABLE:
        st.warning("⚠️ IA non disponible — Configurez votre clé API Google dans `.env`")
        return 0

    try:
        doc = fitz.open(stream=uploaded_pdf.read(), filetype="pdf")
    except Exception as e:
        st.error(f"Erreur ouverture PDF: {e}")
        return 0

    data_found = []
    total_pages = min(page_fin, len(doc)) - page_debut + 1
    progress_bar = st.progress(0, text="Extraction en cours...")

    for i, page_num in enumerate(range(page_debut - 1, min(page_fin, len(doc)))):
        page = doc.load_page(page_num)
        texte = page.get_text("text")

        progress_bar.progress(
            (i + 1) / total_pages,
            text=f"Page {page_num + 1}/{min(page_fin, len(doc))}...",
        )

        if len(texte) > 50:
            resultats = extraire_erreurs_texte(texte)
        else:
            # Page probablement sous forme d'image → OCR via IA
            pix = page.get_pixmap(dpi=200)
            resultats = extraire_erreurs_image(pix.tobytes("png"))

        data_found.extend(resultats)
        time.sleep(1)  # Rate limiting API

    progress_bar.empty()

    # Construire les rows pour la base
    rows_hex = []
    rows_txt = []
    for item in data_found:
        code = str(item.get("Code", "")).strip()
        if len(code) > 1:
            rows_hex.append({
                "Code": code,
                "Message": item.get("Message", ""),
                "Niveau": "ATTENTION",
                "Type": "Manuel PDF",
            })
            rows_txt.append({
                "Mot_Cle": code,
                "Type": "Manuel PDF",
                "Priorite": "MOYENNE",
                "Cause": item.get("Cause", ""),
                "Solution": item.get("Solution", ""),
            })

    if rows_hex:
        ajouter_codes_batch(excel_path, rows_hex, rows_txt)

    return len(rows_hex)


def importer_docx(uploaded_docx, excel_path):
    """
    Importe des codes d'erreur depuis un document Word (.docx).
    Extrait le texte des paragraphes et tableaux, puis utilise l'IA
    pour identifier les codes d'erreur et solutions.
    Retourne le nombre de codes extraits.
    """
    if not AI_AVAILABLE:
        st.warning("⚠️ IA non disponible — Configurez votre clé API Google dans `.env`")
        return 0

    try:
        from docx import Document
    except ImportError:
        st.error("❌ Module `python-docx` non installé. Lancez : `pip install python-docx`")
        return 0

    try:
        doc = Document(uploaded_docx)
    except Exception as e:
        st.error(f"Erreur ouverture DOCX: {e}")
        return 0

    # 1. Extraire tout le texte (paragraphes + tableaux)
    all_text_blocks = []

    # Paragraphes
    current_block = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_block.append(text)
        elif current_block:
            all_text_blocks.append("\n".join(current_block))
            current_block = []
    if current_block:
        all_text_blocks.append("\n".join(current_block))

    # Tableaux (souvent les codes d'erreur sont dans des tableaux)
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_data))
        if table_text:
            all_text_blocks.append("\n".join(table_text))

    if not all_text_blocks:
        st.warning("Le document semble vide.")
        return 0

    # 2. Découper en chunks de ~2000 caractères pour l'IA
    chunks = []
    current_chunk = ""
    for block in all_text_blocks:
        if len(current_chunk) + len(block) > 2000:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = block
        else:
            current_chunk += "\n\n" + block if current_chunk else block
    if current_chunk:
        chunks.append(current_chunk)

    # 3. Envoyer chaque chunk à l'IA pour extraction
    data_found = []
    progress_bar = st.progress(0, text="Analyse du document Word...")

    for i, chunk in enumerate(chunks):
        progress_bar.progress(
            (i + 1) / len(chunks),
            text=f"Analyse bloc {i + 1}/{len(chunks)}..."
        )
        resultats = extraire_erreurs_texte(chunk)
        data_found.extend(resultats)
        time.sleep(1)  # Rate limiting API

    progress_bar.empty()

    # 4. Construire les rows pour la base
    rows_hex = []
    rows_txt = []
    for item in data_found:
        code = str(item.get("Code", "")).strip()
        if len(code) > 1:
            rows_hex.append({
                "Code": code,
                "Message": item.get("Message", ""),
                "Niveau": "ATTENTION",
                "Type": "Manuel DOCX",
            })
            rows_txt.append({
                "Mot_Cle": code,
                "Type": "Manuel DOCX",
                "Priorite": "MOYENNE",
                "Cause": item.get("Cause", ""),
                "Solution": item.get("Solution", ""),
            })

    if rows_hex:
        ajouter_codes_batch(excel_path, rows_hex, rows_txt)

    return len(rows_hex)
